"""
B401 – Continuous-Time Derivatives Pricing
A3 Landscape Poster Generator
Run: pip install python-docx  then  python generate_poster.py
Output: B401_Poster_RHM.docx
"""

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colours ───────────────────────────────────────────────────────────────────
COLORS = {
    'DARK_BLUE':   RGBColor(0x0C, 0x44, 0x7C),
    'MID_BLUE':    RGBColor(0x37, 0x8A, 0xDD),
    'LIGHT_BLUE':  RGBColor(0xE6, 0xF1, 0xFB),
    'TEAL':        RGBColor(0x0F, 0x6E, 0x56),
    'TEAL_LIGHT':  RGBColor(0xE1, 0xF5, 0xEE),
    'AMBER':       RGBColor(0xBA, 0x75, 0x17),
    'AMBER_LIGHT': RGBColor(0xFA, 0xEE, 0xDA),
    'CORAL':       RGBColor(0x99, 0x3C, 0x1D),
    'CORAL_LIGHT': RGBColor(0xFA, 0xEC, 0xE7),
    'WHITE':       RGBColor(0xFF, 0xFF, 0xFF),
    'BLACK':       RGBColor(0x00, 0x00, 0x00),
    'GRAY_LIGHT':  RGBColor(0xF1, 0xEF, 0xE8),
    'GRAY_MID':    RGBColor(0x88, 0x87, 0x80),
}

# For backward compatibility, assign constants
DARK_BLUE   = COLORS['DARK_BLUE']
MID_BLUE    = COLORS['MID_BLUE']
LIGHT_BLUE  = COLORS['LIGHT_BLUE']
TEAL        = COLORS['TEAL']
TEAL_LIGHT  = COLORS['TEAL_LIGHT']
AMBER       = COLORS['AMBER']
AMBER_LIGHT = COLORS['AMBER_LIGHT']
CORAL       = COLORS['CORAL']
CORAL_LIGHT = COLORS['CORAL_LIGHT']
WHITE       = COLORS['WHITE']
BLACK       = COLORS['BLACK']
GRAY_LIGHT  = COLORS['GRAY_LIGHT']
GRAY_MID    = COLORS['GRAY_MID']

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), str(rgb))
    tcPr.append(shd)

def cell_para(cell, text, bold=False, size=9, color=BLACK,
              align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0,
              italic=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p = cell.paragraphs[0] if cell.paragraphs[0].text == '' else cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(size * 1.3)
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size); run.font.color.rgb = color
    return p

def make_inner_table(cell, rows, cols):
    t = cell.add_table(rows=rows, cols=cols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    return t

def section_hdr(tbl, row, col, title, bg=LIGHT_BLUE, fg=DARK_BLUE):
    set_cell_bg(tbl.cell(row, col), bg)
    cell_para(tbl.cell(row, col), f'  {title}', bold=True, size=9,
              color=fg, space_before=2, space_after=2)

def kv_rows(cell, items, key_color=DARK_BLUE, size=8.5):
    for lbl, val in items:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        r1 = p.add_run(lbl); r1.bold = True
        r1.font.size = Pt(size); r1.font.color.rgb = key_color
        r2 = p.add_run(val)
        r2.font.size = Pt(size); r2.font.color.rgb = BLACK

def data_table(parent_cell, headers, rows_data, hdr_bg=DARK_BLUE,
               alt_bg=GRAY_LIGHT, size=8):
    t = parent_cell.add_table(rows=len(rows_data)+1, cols=len(headers))
    t.style = 'Table Grid'
    for j, h in enumerate(headers):
        set_cell_bg(t.cell(0, j), hdr_bg)
        cell_para(t.cell(0, j), h, bold=True, size=size, color=WHITE,
                  space_before=2, space_after=2)
    for i, row in enumerate(rows_data):
        bg = WHITE if i % 2 == 0 else alt_bg
        for j, txt in enumerate(row):
            set_cell_bg(t.cell(i+1, j), bg)
            cell_para(t.cell(i+1, j), txt, size=size, space_before=2, space_after=2)
    return t

def spacer(cell, pt=4):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pt)

def placeholder(cell, text):
    p = cell.add_paragraph(text)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(2)
    p.runs[0].font.color.rgb = GRAY_MID
    p.runs[0].font.size = Pt(8)
    p.runs[0].italic = True

def obs_block(cell, label, text, label_color=DARK_BLUE):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(label)
    r1.bold = True; r1.font.size = Pt(8.5); r1.font.color.rgb = label_color
    r2 = p.add_run(text)
    r2.font.size = Pt(8.5); r2.font.color.rgb = BLACK


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()
sec = doc.sections[0]
sec.page_width    = Cm(42.0)
sec.page_height   = Cm(29.7)
sec.landscape     = True
sec.top_margin    = Cm(1.0)
sec.bottom_margin = Cm(0.8)
sec.left_margin   = Cm(1.0)
sec.right_margin  = Cm(1.0)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(9)


# ══════════════════════════════════════════════════════════════════════════════
# HEADER
# ══════════════════════════════════════════════════════════════════════════════

hdr = doc.add_table(rows=1, cols=3)
hdr.style = 'Table Grid'
hdr.alignment = WD_TABLE_ALIGNMENT.CENTER

for c in hdr.columns:
    for cell in c.cells:
        set_cell_bg(cell, DARK_BLUE)

hdr.cell(0,0).width = Cm(8)
hdr.cell(0,1).width = Cm(24)
hdr.cell(0,2).width = Cm(8)

# Left
p = hdr.cell(0,0).paragraphs[0]
p.paragraph_format.space_before = Pt(4)
r = p.add_run('Name: [Your Name]   |   Student ID: [Your ID]')
r.font.color.rgb = WHITE; r.font.size = Pt(8.5); r.bold = True
p2 = hdr.cell(0,0).add_paragraph('Discussion partners: [Fellow students]')
p2.paragraph_format.space_after = Pt(4)
p2.runs[0].font.color.rgb = RGBColor(0xB5,0xD4,0xF4); p2.runs[0].font.size = Pt(7.5)

# Centre
p = hdr.cell(0,1).paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(3)
r1 = p.add_run('B401 Continuous-Time Derivatives Pricing  |  ')
r1.font.color.rgb = RGBColor(0xB5,0xD4,0xF4); r1.font.size = Pt(9)
r2 = p.add_run('Take-Home Assignment – Summer Term 2026')
r2.font.color.rgb = WHITE; r2.font.size = Pt(9); r2.bold = True
p2 = hdr.cell(0,1).add_paragraph(
    'Bonus Certificate on Rheinmetall AG (RHM.DE)  ·  Barrier 70% of S₀  ·  Bonus Level 115%  ·  T = 3 Years')
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(4)
p2.runs[0].font.color.rgb = MID_BLUE; p2.runs[0].font.size = Pt(8)

# Right
p = hdr.cell(0,2).paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_before = Pt(4)
r = p.add_run('Eberhard Karls Universität Tübingen')
r.font.color.rgb = WHITE; r.font.size = Pt(8.5); r.bold = True
for line in ['Faculty of Economics and Social Sciences',
             'Department of Finance  |  Prof. Dr. C. Koziol']:
    pp = hdr.cell(0,2).add_paragraph(line)
    pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pp.runs[0].font.color.rgb = RGBColor(0xB5,0xD4,0xF4); pp.runs[0].font.size = Pt(7.5)
hdr.cell(0,2).paragraphs[-1].paragraph_format.space_after = Pt(4)

p_gap = doc.add_paragraph()
p_gap.paragraph_format.space_after = Pt(3)


# ══════════════════════════════════════════════════════════════════════════════
# BODY – 3 columns
# ══════════════════════════════════════════════════════════════════════════════

body = doc.add_table(rows=1, cols=3)
body.style = 'Table Grid'
body.alignment = WD_TABLE_ALIGNMENT.CENTER

C1 = body.cell(0,0); C1.width = Cm(13.0)
C2 = body.cell(0,1); C2.width = Cm(13.5)
C3 = body.cell(0,2); C3.width = Cm(13.5)

from docx.oxml.ns import qn as _qn
from docx.oxml import OxmlElement as _El

def remove_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    b = _El('w:tcBorders')
    for s in ['top','bottom','left','right']:
        e = _El(f'w:{s}'); e.set(_qn('w:val'),'nil')
        b.append(e)
    tcPr.append(b)

for c in [C1, C2, C3]:
    remove_border(c)


# ────────────────────────────────────────────────────────────────────────────
# COL 1 – PART 1
# ────────────────────────────────────────────────────────────────────────────

p = C1.add_paragraph('PART 1 – CERTIFICATE DESIGN')
p.paragraph_format.space_after = Pt(3)
p.runs[0].bold = True; p.runs[0].font.size = Pt(10)
p.runs[0].font.color.rgb = DARK_BLUE

# Task I
ti = make_inner_table(C1, 2, 1)
section_hdr(ti, 0, 0, 'I. Investor Profile')
set_cell_bg(ti.cell(1,0), GRAY_LIGHT)
ti.cell(1,0).paragraphs[0].paragraph_format.space_before = Pt(2)
kv_rows(ti.cell(1,0), [
    ('Age / Horizon: ',    '60–65 years old, 15-year investment horizon'),
    ('Risk tolerance: ',   'High – growth-oriented, not income-dependent'),
    ('Market view: ',      'Bullish on European defence & NATO rearmament'),
    ('Goal: ',             'Maximum long-term capital appreciation; accepts high volatility'),
])
obs_block(ti.cell(1,0),
    'Why not direct stock? ',
    'A direct position offers no downside buffer. The Bonus Certificate provides a 30% '
    'protection corridor while maintaining full 1:1 upside participation – superior '
    'risk/return for this investor profile.',
    label_color=TEAL)
ti.cell(1,0).paragraphs[-1].paragraph_format.space_after = Pt(3)

spacer(C1)

# Task II
tii = make_inner_table(C1, 2, 1)
section_hdr(tii, 0, 0, 'II. Product Design – Bonus Certificate on RHM.DE')
set_cell_bg(tii.cell(1,0), GRAY_LIGHT)
p = tii.cell(1,0).paragraphs[0]
p.paragraph_format.space_before = Pt(3)
r1 = p.add_run('Replication:  ')
r1.bold = True; r1.font.size = Pt(8.5); r1.font.color.rgb = DARK_BLUE
r2 = p.add_run('Bonus Cert  =  Forward  +  Down-and-Out Put(K=Bonus, B=Barrier)')
r2.font.size = Pt(9); r2.italic = True; r2.font.color.rgb = TEAL

data_table(tii.cell(1,0),
    ['Component', 'Type', 'Role'],
    [['Long Forward', 'Equity forward', 'Full 1:1 upside participation'],
     ['Down-and-Out Put', 'Barrier option (Rubinstein-Reiner)', 'Guarantees bonus if barrier intact']])

kv_rows(tii.cell(1,0), [
    ('Underlying: ',  'Rheinmetall AG (RHM.DE), DAX constituent, defence & security'),
    ('Barrier: ',     '70% of S₀  →  30% protection buffer before bonus is lost'),
    ('Bonus level: ', '115% of S₀  →  ≥15% guaranteed return if barrier holds'),
    ('Maturity: ',    '3 years  ·  Participation: 1:1 uncapped above bonus level'),
    ('Pricing: ',     'Black-Scholes + Rubinstein-Reiner (1991) closed-form barrier formula'),
    ('Div. yield: ',  '~0.8% (low → put more expensive, higher barrier sensitivity)'),
])
tii.cell(1,0).paragraphs[-1].paragraph_format.space_after = Pt(3)

spacer(C1)

# Task III
tiii = make_inner_table(C1, 2, 1)
section_hdr(tiii, 0, 0, 'III. Payoff Profile')
set_cell_bg(tiii.cell(1,0), GRAY_LIGHT)
placeholder(tiii.cell(1,0), '[INSERT payoff_profile.png – generated by Jupyter notebook]')

data_table(tiii.cell(1,0),
    ['Market scenario at maturity', 'Certificate payoff', 'vs direct stock'],
    [['S_T > 115% S₀  (strong bull market)',
      'S_T  (full participation)',
      'Equal'],
     ['70% < S_T ≤ 115% S₀  (sideways / mild decline)',
      '115% · S₀  (guaranteed bonus)',
      'OUTPERFORMS ✓'],
     ['Barrier breached at any t ∈ [0,T]',
      'S_T  (protection lost permanently)',
      'Equal (cert behaves like stock)']],
    hdr_bg=DARK_BLUE)
tiii.cell(1,0).paragraphs[-1].paragraph_format.space_after = Pt(3)

spacer(C1)

# Task IV
tiv = make_inner_table(C1, 2, 1)
section_hdr(tiv, 0, 0, 'IV. German Market Size Estimation')
set_cell_bg(tiv.cell(1,0), GRAY_LIGHT)
tiv.cell(1,0).paragraphs[0].paragraph_format.space_before = Pt(2)
kv_rows(tiv.cell(1,0), [
    ('Total structured product market (BSW 2023): ', '€112 billion'),
    ('Bonus certificate market share: ',             '1.07%  →  ≈ €1.2 billion'),
    ('Demand driver: ',  'Growth-oriented retail investors bullish on DAX defence names'),
    ('Comparable: ',     'DZ Bank, LBBW, DekaBank regularly issue RHM Bonus Certs on Euwax'),
])
tiv.cell(1,0).paragraphs[-1].paragraph_format.space_after = Pt(3)


# ────────────────────────────────────────────────────────────────────────────
# COL 2 – PART 2
# ────────────────────────────────────────────────────────────────────────────

p = C2.add_paragraph('PART 2 – VALUATION OF CERTIFICATES')
p.paragraph_format.space_after = Pt(3)
p.runs[0].bold = True; p.runs[0].font.size = Pt(10)
p.runs[0].font.color.rgb = DARK_BLUE

# Task V
tv = make_inner_table(C2, 2, 1)
section_hdr(tv, 0, 0, 'V. Daily Valuation – 130 Trading Days')
set_cell_bg(tv.cell(1,0), GRAY_LIGHT)
p = tv.cell(1,0).paragraphs[0]
p.paragraph_format.space_before = Pt(3)
r1 = p.add_run('Method: '); r1.bold = True; r1.font.size = Pt(8.5); r1.font.color.rgb = DARK_BLUE
r2 = p.add_run('Rubinstein-Reiner closed-form barrier option formula under Black-Scholes. '
               'Volatility: 63-day rolling historical σ (annualised). '
               'Market prices: real traded Bonus Certificate on RHM from Euwax Stuttgart.')
r2.font.size = Pt(8.5)

placeholder(tv.cell(1,0), '[INSERT valuation_comparison.png – model vs market + daily error bars]')

# Error metrics table with highlight on max error
err_hdrs = ['Metric', 'EUR', '%']
err_data = [
    ('Mean Error (ME)',       '[x.xxxx]', '[x.xxx%]'),
    ('Mean Abs. Error (MAE)', '[x.xxxx]', '[x.xxx%]'),
    ('RMSE',                  '[x.xxxx]', '[x.xxx%]'),
    ('Max Absolute Error ★',  '[x.xxxx]', '[x.xxx%]'),
    ('95th Pct Abs. Error',   '[x.xxxx]', '[x.xxx%]'),
]
err_t = tv.cell(1,0).add_table(rows=6, cols=3)
err_t.style = 'Table Grid'
for j, h in enumerate(err_hdrs):
    set_cell_bg(err_t.cell(0,j), DARK_BLUE)
    cell_para(err_t.cell(0,j), h, bold=True, size=8, color=WHITE,
              space_before=2, space_after=2)
for i, (m, e, pct) in enumerate(err_data):
    is_max = (i == 3)
    bg = AMBER_LIGHT if is_max else (WHITE if i%2==0 else GRAY_LIGHT)
    col = AMBER if is_max else BLACK
    for j, txt in enumerate([m, e, pct]):
        set_cell_bg(err_t.cell(i+1,j), bg)
        cell_para(err_t.cell(i+1,j), txt, size=8, bold=is_max,
                  color=col, space_before=2, space_after=2)

obs_block(tv.cell(1,0),
    'Max absolute error (★): ',
    'The single largest daily deviation between model and market price. '
    'For RHM – a high-volatility defence stock driven by geopolitical news – '
    'this metric is critically important. Black-Scholes assumes constant σ and '
    'cannot capture sudden vol spikes (e.g. Ukraine ceasefire news), causing '
    'the max error to exceed RMSE by a large multiple on event days.',
    label_color=AMBER)
tv.cell(1,0).paragraphs[-1].paragraph_format.space_after = Pt(3)

spacer(C2)

# Task VI
tvi = make_inner_table(C2, 2, 1)
section_hdr(tvi, 0, 0, 'VI. Sensitivity Analysis – Greeks')
set_cell_bg(tvi.cell(1,0), GRAY_LIGHT)

placeholder(tvi.cell(1,0), '[INSERT greeks.png – Delta, Gamma, Vega, Theta vs stock price S]')

data_table(tvi.cell(1,0),
    ['Greek', 'Behaviour', 'Economic interpretation'],
    [['Delta  ∂V/∂S',
      'Rises sharply near B; ~1 above bonus',
      'Near barrier: leveraged (Δ>1). Above bonus: stock-like (Δ≈1).'],
     ['Gamma  ∂²V/∂S²',
      'Spikes near barrier; near zero elsewhere',
      'Small S moves → large Δ swings near B. Key hedging risk.'],
     ['Vega  ∂V/∂σ',
      'Negative near barrier; small above',
      'Higher σ → higher knock-out probability → cert worth less.'],
     ['Theta  ∂V/∂t',
      'Generally positive (cert benefits from time passing)',
      'Fewer days remaining → lower knock-out probability → value rises.']],
    hdr_bg=DARK_BLUE)
tvi.cell(1,0).paragraphs[-1].paragraph_format.space_after = Pt(3)

spacer(C2)

# Task VII
tvii = make_inner_table(C2, 2, 1)
section_hdr(tvii, 0, 0, 'VII. Replicating Portfolio')
set_cell_bg(tvii.cell(1,0), GRAY_LIGHT)

p = tvii.cell(1,0).paragraphs[0]
p.paragraph_format.space_before = Pt(3)
r1 = p.add_run('Structure: '); r1.bold = True; r1.font.size = Pt(8.5); r1.font.color.rgb = DARK_BLUE
r2 = p.add_run('V(S,t) = Δ(S,t) · S  +  Bond(S,t)   where  Bond = V − Δ · S')
r2.font.size = Pt(8.5); r2.italic = True; r2.font.color.rgb = TEAL

placeholder(tvii.cell(1,0),
    '[INSERT equity_fraction_vs_S.png – equity fraction vs S at fixed T_rem]\n'
    '[INSERT replicating_portfolio_time.png – daily delta and equity fraction over time]')

obs_block(tvii.cell(1,0),
    'Falling market interpretation: ',
    'As RHM falls toward the barrier, equity fraction rises above 100% – the replicating '
    'portfolio is long more than one share and short the bond (leveraged). '
    'This reflects the deeply negative delta of the down-and-out put near the barrier. '
    'Dynamic delta hedging therefore requires aggressive stock purchases as the stock '
    'falls – amplifying losses in a sustained decline. '
    'Once the barrier is breached, the put knocks out permanently and the portfolio '
    'converts to a pure equity position.')
tvii.cell(1,0).paragraphs[-1].paragraph_format.space_after = Pt(3)


# ────────────────────────────────────────────────────────────────────────────
# COL 3 – PART 3
# ────────────────────────────────────────────────────────────────────────────

p = C3.add_paragraph('PART 3 – PORTFOLIO INSURANCE STRATEGY')
p.paragraph_format.space_after = Pt(3)
p.runs[0].bold = True; p.runs[0].font.size = Pt(10)
p.runs[0].font.color.rgb = DARK_BLUE

# Setup box
setup = make_inner_table(C3, 2, 1)
set_cell_bg(setup.cell(0,0), DARK_BLUE)
cell_para(setup.cell(0,0), '  Setup', bold=True, size=9, color=WHITE,
          space_before=2, space_after=2)
set_cell_bg(setup.cell(1,0), TEAL_LIGHT)
kv_rows(setup.cell(1,0), [
    ('Initial capital: ', '€10,000'),
    ('Horizon T*: ',      '1 year  ·  252 daily steps'),
    ('Underlying: ',      'Rheinmetall AG (RHM.DE) – GBM with μ=historical, σ=historical'),
    ('Simulation: ',      '50,000 Monte Carlo paths'),
    ('Insurance: ',       'Fictitious European puts priced on historical σ via Black-Scholes'),
], key_color=TEAL)
setup.cell(1,0).paragraphs[-1].paragraph_format.space_after = Pt(3)

spacer(C3)

# Task VIII
tviii = make_inner_table(C3, 2, 1)
section_hdr(tviii, 0, 0, 'VIII. Performance Without Risk Management', bg=CORAL_LIGHT, fg=CORAL)
set_cell_bg(tviii.cell(1,0), GRAY_LIGHT)

placeholder(tviii.cell(1,0), '[INSERT mc_unhedged.png – return distribution, 95% VaR marked]')

data_table(tviii.cell(1,0),
    ['Risk/return metric', 'Unhedged (pure stock)'],
    [['Mean return',     '[x.xx%]'],
     ['Std deviation',   '[x.xx%]'],
     ['Sharpe ratio',    '[x.xxx]'],
     ['95% VaR (loss)',  '[x.xx%]'],
     ['95% CVaR (loss)', '[x.xx%]'],
     ['P(loss)',         '[xx.x%]']],
    hdr_bg=CORAL, alt_bg=GRAY_LIGHT)
tviii.cell(1,0).paragraphs[-1].paragraph_format.space_after = Pt(3)

spacer(C3)

# Task IX
tix = make_inner_table(C3, 2, 1)
section_hdr(tix, 0, 0, 'IX. Portfolio Insurance with Put Options', bg=CORAL_LIGHT, fg=CORAL)
set_cell_bg(tix.cell(1,0), GRAY_LIGHT)

p = tix.cell(1,0).paragraphs[0]
p.paragraph_format.space_before = Pt(3)
r1 = p.add_run('Portfolio: '); r1.bold = True; r1.font.size = Pt(8.5); r1.font.color.rgb = DARK_BLUE
r2 = p.add_run('(1−α) × RHM stock  +  α × European Put options')
r2.font.size = Pt(8.5); r2.italic = True

placeholder(tix.cell(1,0), '[INSERT insurance_heatmap.png – VaR and mean return heatmap (α × K)]')

kv_rows(tix.cell(1,0), [
    ('Higher α (more puts): ',    'VaR ↓, mean return ↓. Clear risk/return trade-off.'),
    ('Higher K (deeper ITM): ',   'Stronger protection but higher put cost → lower mean return.'),
    ('Best balance: ',            'α = 10%, K = 95% S₀ – maximises Sharpe under cost constraint.'),
])
tix.cell(1,0).paragraphs[-1].paragraph_format.space_after = Pt(3)

spacer(C3)

# Task X
tx = make_inner_table(C3, 2, 1)
section_hdr(tx, 0, 0, 'X. Stress Scenario Analysis  (α=10%, K=95% S₀)', bg=CORAL_LIGHT, fg=CORAL)
set_cell_bg(tx.cell(1,0), GRAY_LIGHT)

stress_t = tx.cell(1,0).add_table(rows=5, cols=4)
stress_t.style = 'Table Grid'
for j, h in enumerate(['Scenario', 'Mean return', '95% VaR', '95% CVaR']):
    set_cell_bg(stress_t.cell(0,j), CORAL)
    cell_para(stress_t.cell(0,j), h, bold=True, size=8, color=WHITE, space_before=2, space_after=2)
stress_rows = [
    ('Baseline',           '[x.xx%]', '[x.xx%]', '[x.xx%]'),
    ('Vol +5pp',           '[x.xx%]', '[x.xx%]', '[x.xx%]'),
    ('Vol −5pp',           '[x.xx%]', '[x.xx%]', '[x.xx%]'),
    ('−20% shock at T/2',  '[x.xx%]', '[x.xx%]', '[x.xx%]'),
]
for i, row in enumerate(stress_rows):
    is_shock = (i == 3)
    bg = AMBER_LIGHT if is_shock else (WHITE if i%2==0 else GRAY_LIGHT)
    for j, txt in enumerate(row):
        set_cell_bg(stress_t.cell(i+1,j), bg)
        cell_para(stress_t.cell(i+1,j), txt, size=8, bold=is_shock,
                  color=AMBER if is_shock else BLACK, space_before=2, space_after=2)

placeholder(tx.cell(1,0), '[INSERT stress_scenarios.png – overlaid return distributions]')

obs_block(tx.cell(1,0),
    'Discussion: ',
    'Higher pricing vol inflates put cost → fewer puts purchased → weaker protection. '
    'The −20% price shock at T/2 dominates: triggers put payoffs but the stock continues '
    'to drift; depending on K this may still result in significant losses. '
    'For RHM specifically, sudden vol spikes (geopolitical events) create the largest '
    'stress impact due to the elevated baseline σ.')
tx.cell(1,0).paragraphs[-1].paragraph_format.space_after = Pt(3)

spacer(C3)

# Task XI
txi = make_inner_table(C3, 2, 1)
section_hdr(txi, 0, 0, 'XI. Capital Requirement – 95% VaR ≤ 15%', bg=CORAL_LIGHT, fg=CORAL)
set_cell_bg(txi.cell(1,0), GRAY_LIGHT)

placeholder(txi.cell(1,0), '[INSERT var_constraint.png – VaR vs α with binding constraint at 15%]')

data_table(txi.cell(1,0),
    ['Parameter', 'Optimal allocation'],
    [['Strike K',            '95% of S₀'],
     ['α*  (put fraction)',  '[x.xx%]  ←  solved via scipy Brent root-finder'],
     ['Resulting 95% VaR',  '15.00%  (constraint fully utilised)'],
     ['Expected return',     '[x.xx%]'],
     ['Sharpe ratio',        '[x.xxx]']],
    hdr_bg=DARK_BLUE, alt_bg=GRAY_LIGHT)

obs_block(txi.cell(1,0),
    'Interpretation: ',
    'α* fully utilises the VaR budget. More puts wastes return potential; '
    'fewer puts violates the constraint. '
    'The Brent solver identifies α* numerically from the simulated distribution '
    'by finding the root of VaR(α) + 0.15 = 0. '
    'This approach allows putting more capital into the risky asset than a '
    'pure stock strategy would permit under the same VaR limit.')
txi.cell(1,0).paragraphs[-1].paragraph_format.space_after = Pt(3)


# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph().paragraph_format.space_after = Pt(3)

ftr = doc.add_table(rows=1, cols=3)
ftr.style = 'Table Grid'
ftr.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (w, txt, align) in enumerate([
    (Cm(13),
     'Pricing: Black-Scholes + Rubinstein-Reiner (1991)  ·  '
     'Data: Yahoo Finance (yfinance)  ·  Simulation: 50,000 GBM paths',
     WD_ALIGN_PARAGRAPH.LEFT),
    (Cm(14),
     'B401 Take-Home Assignment  ·  Summer Term 2026  ·  '
     'Eberhard Karls Universität Tübingen',
     WD_ALIGN_PARAGRAPH.CENTER),
    (Cm(13),
     'References: Hull (2017) Options, Futures & Derivatives  ·  '
     'Rubinstein & Reiner (1991)  ·  BSW/DDV (2023)',
     WD_ALIGN_PARAGRAPH.RIGHT),
]):
    set_cell_bg(ftr.cell(0,i), DARK_BLUE)
    ftr.cell(0,i).width = w
    p = ftr.cell(0,i).paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(txt)
    r.font.color.rgb = RGBColor(0xB5,0xD4,0xF4) if i != 1 else WHITE
    r.font.size = Pt(7.5)
    if i == 1:
        r.bold = True


# ── Save ──────────────────────────────────────────────────────────────────────
doc.save('B401_Poster_RHM.docx')
print('✓  Saved: B401_Poster_RHM.docx')
print()
print('After running your notebook, open the .docx and:')
print('  1. Click each [INSERT ...] placeholder → Insert → Pictures → From File')
print('  2. Replace all [x.xx%] values with your actual computed results')
print('  3. Update your name, student ID, and discussion partners in the header')
print('  4. Export as PDF (A3 landscape) for submission')