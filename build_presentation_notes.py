"""Build a clean Word document with the full presentation notes."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)

# ── Style helpers ─────────────────────────────────────────────────────────────
def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)  # dark red
    return p

def subheading(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # dark blue
    return p

def sub3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x37, 0x5A, 0x7F)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    return p

def code_para(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(2)
    # light grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p

def bold_body(label, text):
    p = doc.add_paragraph()
    run = p.add_run(label + ' ')
    run.bold = True
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    return p

def divider():
    doc.add_paragraph('─' * 80)

def qa_block(question, answer):
    p = doc.add_paragraph()
    r = p.add_run('Q: ' + question)
    r.bold = True
    r.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
    p.paragraph_format.space_before = Pt(6)

    p2 = doc.add_paragraph()
    r2 = p2.add_run('A: ')
    r2.bold = True
    p2.add_run(answer)
    p2.paragraph_format.left_indent = Inches(0.2)
    p2.paragraph_format.space_after = Pt(6)

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('B401 — Continuous-Time Derivatives Pricing')
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('Presentation Notes & Grader Q&A Preparation')
r2.font.size = Pt(13)
r2.bold = True

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('University of Tübingen · Prof. Dr. Christian Koziol · Summer Term 2026\n'
             'Student: Daanish Muzaffar (7259472)')

doc.add_paragraph()
grade_p = doc.add_paragraph()
grade_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
gr = grade_p.add_run('Grade: 99 / 115 pts  →  1.7 (Gut)')
gr.bold = True
gr.font.size = Pt(14)
gr.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 0 — GRADE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════
heading('Grade Summary')

# Table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Part'
hdr[1].text = 'Task'
hdr[2].text = 'Max'
hdr[3].text = 'Score'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows_data = [
    ('Part 1', 'I — Investor Profile', '10', '8'),
    ('Part 1', 'II — Product Design', '10', '8'),
    ('Part 1', 'III — Product Payoff', '10', '9'),
    ('Part 1', 'IV — Market Size', '10', '8'),
    ('Part 2', 'V — Valuation', '15', '14'),
    ('Part 2', 'VI — Greeks', '12', '10'),
    ('Part 2', 'VII — Replicating Portfolio', '10', '7'),
    ('Part 3', 'VIII — MC (no insurance)', '8', '8'),
    ('Part 3', 'IX — MC (with insurance)', '10', '9'),
    ('Part 3', 'X — Stress Scenarios', '7', '7'),
    ('Part 3', 'XI — VaR Constraint', '8', '7'),
    ('Quality', 'Code', '5', '4'),
    ('TOTAL', '', '115', '99'),
]
for rd in rows_data:
    row = table.add_row().cells
    for i, val in enumerate(rd):
        row[i].text = val
    if rd[0] == 'TOTAL':
        for cell in row:
            for run in cell.paragraphs[0].runs:
                run.bold = True

doc.add_paragraph()
body('Grading thresholds: 104–115 → 1.0 (Sehr gut) | 92–103 → 1.3–1.7 (Sehr gut/Gut) | '
     '81–91 → 2.0–2.3 (Gut). Score of 99 falls in the 92–103 band, at 99/103 through the band → 1.7.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — SETUP
# ═══════════════════════════════════════════════════════════════════════════════
heading('Setup — What Happens Before Any Task (Cells 1–14)')

sub3('Libraries imported')
bullet('numpy, pandas — numerical operations and time-series handling')
bullet('scipy.stats.norm — Black-Scholes closed-form analytics')
bullet('scipy.optimize.brentq — root-finding for Task XI')
bullet('yfinance — live RHM.DE price download')
bullet('matplotlib — all charts')

sub3('1. Data Download (Cell 7)')
body('Downloads Rheinmetall AG (RHM.DE) dividend- and split-adjusted daily closes '
     'from Yahoo Finance (Jan 2023 – Jun 2026). Computes daily log-returns:')
code_para('r_t = ln(S_t / S_{t-1})')

sub3('2. EWMA Volatility (Cell 7) — RiskMetrics λ = 0.94')
body('Exponentially Weighted Moving Average variance, updated recursively each day:')
code_para('σ²_t = λ·σ²_{t-1} + (1−λ)·r²_t     where λ = 0.94')
body('λ = 0.94 is the J.P. Morgan RiskMetrics industry standard for daily returns. '
     'Each return observation receives weight (1−0.94) = 6%, decaying exponentially. '
     'This EWMA vol is the day-by-day pricing input for Parts 1 and 2.')

sub3('3. Svensson (1994) Yield Curve (Cell 13)')
body('ECB AAA sovereign curve parameters (β₀ β₁ β₂ β₃ τ₁ τ₂) loaded from a local CSV '
     '(ECB Data Portal download, no runtime API calls). Formula:')
code_para('r(t) = β₀ + β₁·f(t/τ₁) + β₂·[f(t/τ₁)−e^{−t/τ₁}] + β₃·[f(t/τ₂)−e^{−t/τ₂}]')
code_para('f(x) = (1 − e^{−x}) / x')
body('get_risk_free_rate(date, T_rem) forward-fills on ECB business days and returns '
     'the maturity-matched spot rate for each valuation day.')

sub3('4. Global Parameter Registry (Cell 10)')
body('All product constants (K = 2000 EUR, B = 1050 EUR, MATURITY = 18 Jun 2027, '
     'DIV_YIELD from live RHM dividend, EWMA vol) defined once. Every downstream cell reads from here.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — PART 1
# ═══════════════════════════════════════════════════════════════════════════════
heading('PART 1 — Designing the Certificate')

subheading('Task I — Investor Profile  (8/10)')
bold_body('Target investor:', 'German retail investor, 40–60 years old, approaching '
          'pre-retirement. Bullish on European defence rearmament (NATO 2% GDP target, '
          'Sondervermögen, RHM order backlog). Wants upside on Rheinmetall but finds '
          'direct stock (σ ≈ 54% p.a.) unacceptable.')
bold_body('Why a structured product:', 'Direct stock: full downside. Defence ETF: '
          'dilutes theme. Capital-protected note: defence upside with a floor — perfect match.')
bold_body('Five dimensions covered:', 'Age/horizon, risk tolerance, income vs. growth '
          'orientation, market expectations, comparison to alternatives.')
bold_body('Deductions (−2):', 'No quantitative matching of risk tolerance to product vol. '
          'No discussion of reinvestment risk at maturity.')

subheading('Task II — Product Design  (8/10)')
body('Product: Capped Capital Protected Participation Note (CPN) on RHM.DE — three components:')
table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Table Grid'
h2 = table2.rows[0].cells
h2[0].text = 'Component'; h2[1].text = 'Position'; h2[2].text = 'Role'
for c in h2:
    for r in c.paragraphs[0].runs: r.bold = True
for row_data in [
    ('Zero-coupon bond (ZCB)', 'Long', 'Guarantees capital at maturity'),
    ('ATM call (K = S₀)', 'Long × α', 'Captures upside'),
    ('OTM cap call (K = 1.30·S₀)', 'Short × α', 'Finances higher participation rate α'),
]:
    rw = table2.add_row().cells
    for i, v in enumerate(row_data): rw[i].text = v

doc.add_paragraph()
bold_body('Payoff:', 'max(S₀,  S₀ + α·(min(S_T, 1.30·S₀) − S₀))')
bold_body('How α is derived (Cell 18):', 'The product issues at par (price = S₀ = 100). '
          'Setting price = nominal and solving:')
code_para('S₀  =  ZCB(r,T)  +  α × [Call(K=S₀) − Call(K=1.30·S₀)]')
code_para('α   =  (S₀ − ZCB₀) / [Call_ATM − Call_Cap]')
body('Both calls priced with Black-Scholes using Svensson 3-year rate and EWMA vol. '
     'ZCB₀ ≈ 1115.79 EUR, spread ≈ 97.61 EUR → α = 93.4%.')
bold_body('Deductions (−2):', 'No sensitivity of α to rate/vol changes. '
          'Cap at 130% asserted, not optimised.')

subheading('Task III — Payoff Profile  (9/10)')
body('Payoff plotted over S_T ∈ [40%, 190%] of S₀ with three labeled regions:')
bullet('Bearish (S_T < S₀): flat at S₀ — capital protection floor')
bullet('Moderately bullish (S₀ ≤ S_T ≤ 1.30·S₀): rises at slope α = 93.4%')
bullet('Strongly bullish (S_T > 1.30·S₀): flat at cap = S₀ + α·0.30·S₀')
bullet('Green shading = product outperforms direct stock (below S₀ and moderate upside)')
bullet('Blue shading = product underperforms (above cap)')
bold_body('Deduction (−1):', 'Could state GBM-implied probability of each region '
          '(e.g., P(S_T < K) = 8.7% from Task VIII).')

subheading('Task IV — Market Size  (8/10)')
body('Top-down from DDV/BSW 2023 data (slide 13):')
bullet('Total German structured product market: EUR 112bn')
bullet('Kapitalschutzprodukte (KSP) segment: 3.80% = EUR 4.25bn')
bullet('Participation notes within KSP: ~55% = EUR 2.34bn')
bullet('~130 active issues at avg EUR 33mn per issue')
bold_body('Deductions (−2):', 'No separation of addressable market for single-name '
          'product. No demand-side validation.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — PART 2
# ═══════════════════════════════════════════════════════════════════════════════
heading('PART 2 — Valuation of the DU2076 Bonus Cap Certificate')

sub3('The Product — DZ BANK DU2076')
bullet('Underlying: Rheinmetall AG (RHM.DE)')
bullet('Barrier B = 1,050 EUR — continuous, monitored 04 Sep 2025 to 18 Jun 2027')
bullet('Bonus / Cap K = 2,000 EUR')
bullet('Maturity: 18 Jun 2027')
bullet('Payoff if barrier never touched: min(S_T, 2000) + max(2000 − S_T, 0) = always at least 2000')
bullet('Payoff if barrier breached: min(S_T, 2000) — bonus disappears')
bold_body('Decomposition:', 'V_cert = V_capped_forward + V_down-and-out-put')
body('The capped forward pays min(S_T, K) regardless of barrier. '
     'The down-and-out put pays max(K − S_T, 0) only if barrier was never touched. '
     'Together they give the bonus structure.')

subheading('Task V — Daily Valuation  (14/15)')
sub3('The Pricing Engine — CRR Binomial Tree (N = 200)')
body('Cox-Ross-Rubinstein (1979) model. Time discretized into N = 200 equal steps:')
code_para('Δt = T/N')
code_para('u = e^{σ√Δt},   d = 1/u     (up and down factors)')
code_para('p = (e^{(r−q)·Δt} − d) / (u − d)    (risk-neutral probability)')
body('Both components (DO-put and capped forward) are rolled back simultaneously '
     'in a single backward-induction pass. At each node, barrier condition kills '
     'the DO-put: if S_node ≤ 1050 EUR → V_do = 0. Terminal nodes use payoff formulas directly.')
bold_body('Why N = 200:', 'More steps = finer barrier monitoring ≈ continuous. '
          'At N = 200, discretization error is negligible. Tradeoff: ~1 second per tree call.')

sub3('Daily Valuation Loop (Cell 37) — 183 trading days')
body('For each trading day from 04 Sep 2025 to 01 Jun 2026:')
bullet('Compute T_rem = (18 Jun 2027 − date) / 365  (remaining maturity in years)')
bullet('Look up EWMA vol for that day (clipped to 10–60% range)')
bullet('Look up Svensson spot rate r(T_rem) from daily ECB parameters')
bullet('Run N=200 binomial tree → model price')

sub3('Error Metrics (Cell 40)')
table3 = doc.add_table(rows=1, cols=2)
table3.style = 'Table Grid'
h3 = table3.rows[0].cells; h3[0].text = 'Metric'; h3[1].text = 'Value'
for c in h3:
    for r in c.paragraphs[0].runs: r.bold = True
for rd in [
    ('Mean Error (ME)', '−85.54 EUR  (systematic underpricing)'),
    ('MAE', '94.49 EUR'),
    ('RMSE', '108.92 EUR'),
    ('MAPE', '6.11%'),
]:
    rw = table3.add_row().cells; rw[0].text = rd[0]; rw[1].text = rd[1]

doc.add_paragraph()
bold_body('Why model underprice:', 'Market uses implied vol (includes risk premium, '
          'supply/demand). Model uses realized (EWMA) vol. Issuer margin also contributes '
          'to the gap.')

sub3('Implied Vol Back-Solver (Novel Addition — Cell 40)')
body('Using brentq on a faster N=50 tree, implied vol is back-solved for all 183 days:')
bullet('Mean implied vol = 37.77%  vs  EWMA = 44.27%')
bullet('IV model error: ME = −8.89 EUR, MAPE = 0.64% — near-perfect fit')
bullet('Confirms: pricing gap is entirely attributable to the vol assumption, not the tree structure')
bold_body('Deduction (−1):', 'Error quantiles (25th, 75th, 95th percentile) not computed.')

subheading('Task VI — Greeks / Sensitivity Analysis  (10/12)')
sub3('Method — Central Finite Differences (Cell 45)')
code_para('Δ ≈ [V(S+h) − V(S−h)] / (2h·S)          h = 1% stock bump')
code_para('Γ ≈ [V(S+h) − 2V(S) + V(S−h)] / (hS)²')
code_para('V ≈ [V(σ+h) − V(σ−h)] / (2h·100)         h = 2% vol bump')
code_para('Θ ≈ [V(T−1day) − V(T)] / (1/252) / 365')
body('2% vol bump for Vega is deliberate — binomial tree has discrete grid noise '
     'at sub-1% vol perturbations.')

sub3('Reference parameters')
bullet('T = 1.79 years (first DU2076 trading day, 04 Sep 2025)')
bullet('σ = 44.6% (median EWMA vol over observation window)')
bullet('S_ref = 1,708 EUR (first-day spot)')

sub3('Qualitative shapes and interpretation')
bold_body('Delta:', 'Near B=1050 → rises steeply toward 1 (barrier risk). '
          'Above K=2000 → falls toward 0 (capped, stops tracking stock). '
          'Opposite pattern to a vanilla call.')
bold_body('Gamma:', 'Negative near cap (cert is implicitly short an OTM call, '
          'Delta falls as S rises). Spikes positive near barrier (Delta rising steeply).')
bold_body('Vega (NEGATIVE):', 'Higher vol → higher probability of barrier breach → '
          'destroys DO-put value → lower cert value. Opposite to a long vanilla option. '
          'This is the key structural insight for a barrier product.')
bold_body('Theta (POSITIVE):', 'As time passes, less remaining life → lower barrier-breach '
          'probability → bonus more likely to survive → cert value increases. '
          'Opposite to a vanilla option.')
bold_body('Deductions (−2):', 'Sawtooth oscillations from binomial tree discretization '
          'not acknowledged in markdown. Gamma at reference prints as -0.000000 (rounding).')

subheading('Task VII — Replicating Portfolio  (7/10)')
sub3('Theory (Cell 48)')
code_para('V(t) = Δ(t)·S(t) + B(t)')
code_para('Bond: B(t) = V(t) − Δ(t)·S(t)')
code_para('Equity fraction: eq_frac = Δ(t)·S(t) / V(t)')
body('For each of the 183 days: compute Delta via finite differences → derive bond position.')

sub3('Equity Fraction vs S (Static Snapshot — Cell 53)')
body('Sweep S from 0.55·B to 1.55·K at fixed T = 1.78y and median EWMA vol:')
bullet('Near barrier (S → 1050): equity fraction spikes ABOVE 100% — implicit leverage '
       'because Δ is very high but cert value is compressed')
bullet('Between B and K: fraction declines toward 0 as cert approaches cap')
bullet('Above K: fraction near 0, cert is "bond-like"')
bold_body('Falling markets:', 'As RHM falls toward B=1050, Delta rises sharply '
          '(portfolio must hold more stock), equity fraction exceeds 1. '
          'If barrier is breached, bonus disappears — cert converts to capped forward.')
bold_body('Deductions (−3):', 'Daily time-series plot (replicating_portfolio_time.png) '
          'was deleted before submission — a required visualization. '
          'Equity fraction vs S chart is jagged due to binomial tree oscillation.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — PART 3
# ═══════════════════════════════════════════════════════════════════════════════
heading('PART 3 — Portfolio Insurance Strategy')

sub3('Setup (Cell 57)')
bullet('W₀ = EUR 10,000 | Horizon T* = 1 year | Underlying: RHM.DE')
bullet('N_paths = 50,000 | N_steps = 252 (daily frequency) | Seed = 2026')
bullet('μ_hist = 54.05% p.a. (annualised mean of log-returns, 866-day history)')
bullet('σ = 39.88% p.a. (annualised std of log-returns)')
bold_body('GBM path generation:', '')
code_para('S_{t+Δt} = S_t · exp[(μ − ½σ²)·Δt + σ·√Δt·Z]    Z ~ N(0,1)')
body('The Itô correction −½σ² is included because μ_hist is the arithmetic log-price drift '
     '(mean of log-returns already in log space). Seed fixed to 2026 for reproducibility. '
     'Note: Bug B1 (double Itô correction) was fixed before submission.')

subheading('Task VIII — Performance Without Insurance  (8/8 — Full marks)')
body('Apply performance_stats to 50,000 terminal log-returns ln(S_T/S₀):')
table4 = doc.add_table(rows=1, cols=2)
table4.style = 'Table Grid'
h4 = table4.rows[0].cells; h4[0].text = 'Measure'; h4[1].text = 'Value'
for c in h4:
    for r in c.paragraphs[0].runs: r.bold = True
for rd in [
    ('Mean log-return', '54.03%'),
    ('Median log-return', '54.17%'),
    ('Std deviation', '39.92%'),
    ('Skewness', '0.007'),
    ('Sharpe ratio', '1.277'),
    ('95% VaR (loss)', '11.32%'),
    ('95% CVaR (loss)', '28.38%'),
    ('P(loss)', '8.7%'),
]:
    rw = table4.add_row().cells; rw[0].text = rd[0]; rw[1].text = rd[1]

doc.add_paragraph()
bold_body('VaR definition:', '5th percentile of the log-return distribution. '
          '95% VaR = 11.32% means: in the worst 5% of scenarios, loss exceeds 11.32%.')
body('GBM limitations discussed: constant σ (empirical range 20–70%), no jumps '
     '(geopolitical shocks), log-normal tails thinner than empirical.')

subheading('Task IX — Performance With Insurance  (9/10)')
sub3('Method (Cell 64)')
body('Portfolio = (1−α)·stock + α·puts. For each of 16 (α, K) combinations:')
bullet('Price a European put at issuance using Black-Scholes with historical σ = 39.88%')
bullet('n_shares = (1−α)·W₀/S₀,   n_puts = α·W₀/put_price')
bullet('Terminal wealth = stock leg + put payoffs')
bullet('Compute log-return on terminal wealth')
bullet('Grid: α ∈ {5%, 10%, 15%, 20%}  ×  K ∈ {90%, 95%, 100%, 105%}·S₀')

sub3('The Counterintuitive Result — VaR INCREASES with more insurance')
body('Under μ = 54.05%, P(S_T > S₀) = 91.3%. At the 5th-percentile scenario '
     '(a moderate down year), puts are almost always out-of-the-money. '
     'The investor paid the full put premium upfront — this drag on the stock '
     'position is larger than any put payoff received. So the 5th-percentile '
     'return is WORSE with insurance than without.')
body('Exception: K = 105% (ITM puts) — even in mildly bad scenarios, the put partially '
     'pays off and partly offsets the premium cost.')
bold_body('Deduction (−1):', 'Only 4×4 grid; denser α resolution would better trace '
          'the VaR-minimising surface.')

subheading('Task X — Stress Scenarios  (7/7 — Full marks)')
body('Chosen allocation: α = 10%, K = 95%·S₀ (Cell 70)')
table5 = doc.add_table(rows=1, cols=5)
table5.style = 'Table Grid'
h5 = table5.rows[0].cells
for i, t in enumerate(['Scenario', 'Mean', 'VaR 95%', 'CVaR 95%', 'P(loss)']): h5[i].text = t
for c in h5:
    for r in c.paragraphs[0].runs: r.bold = True
for rd in [
    ('Baseline', '44.66%', '16.00%', '16.79%', '13.7%'),
    ('Vol +5pp', '44.52%', '16.77%', '19.40%', '13.7%'),
    ('Vol −5pp', '44.84%', '13.62%', '14.73%', '13.7%'),
    ('−20% at T/2', '24.82%', '17.07%', '17.70%', '29.8%'),
]:
    rw = table5.add_row().cells
    for i, v in enumerate(rd): rw[i].text = v

doc.add_paragraph()
body('Vol stress magnitude (±5pp) from FRTB/Basel IV §21.83. '
     'Price shock (−20%) at T/2 applied by scaling all paths by 0.80 at step 126, '
     'consistent with PRIIPs KID stress methodology (EU Regulation 1286/2014).')
bold_body('Key insight:', 'Baseline VaR = 16.00% already exceeds the Task XI limit of 15%. '
          'This allocation is non-compliant before any stress. Compliant allocation: α* = 7.90%.')

subheading('Task XI — VaR-Constrained Optimal Allocation  (7/8)')
sub3('Method (Cell 74) — scipy.optimize.brentq')
body('Find α* such that 95% VaR(α*) = 15.00% exactly. Steps:')
bullet('Scan α ∈ [0.1%, 40%] at 80 points to confirm VaR is monotone increasing '
       '(required precondition for brentq)')
bullet('Solve: brentq(λ a: VaR(a) − (−0.15), 0.001, 0.40) → convergence guaranteed '
       'because VaR(0.001) < 0.15 and VaR(0.40) > 0.15')
bold_body('Result:', 'α* = 7.90%,  VaR = 15.00%,  E[return] = 46.73%,  Sharpe = 1.148')
body('α* is a regulatory CEILING — the maximum permitted put allocation. '
     'Any more insurance worsens VaR (for this high-drift stock). '
     'Optimal strategy: allocate exactly 7.90% to puts, 92.1% to stock.')
bold_body('Deduction (−1):', 'K = 95% held fixed. A joint optimisation over (α, K) '
          'might find higher expected return at the same binding VaR.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — GRADER Q&A
# ═══════════════════════════════════════════════════════════════════════════════
heading('Anticipated Grader Questions — and How to Answer')

qa_block(
    'Why did you choose RHM.DE as the underlying in both Parts 1 and 2?',
    'For consistency. The same underlying means the EWMA vol estimated for Part 1 design '
    'directly carries into Part 2 pricing. It also allows me to compare the designed '
    'product\'s payoff to the actually-traded DZ BANK certificate on the same stock.'
)

qa_block(
    'Why is your model price systematically below the market price?',
    'The model uses realized (EWMA) volatility. The market prices using implied volatility, '
    'which incorporates a risk premium and supply-demand effects. The implied vol back-solver '
    'confirms: mean IV = 37.77% vs EWMA = 44.27%. Since EWMA is actually higher, the model '
    'should overprice — but the ME is negative (model < market), meaning the issuer\'s margin '
    'is the dominant effect: the market prices the certificate more expensively than a fair '
    'theoretical value because the issuer embeds a spread.'
)

qa_block(
    'Why does VaR INCREASE when you add put insurance?',
    'Because of the high historical drift (54.05% p.a.), only 8.7% of paths end below S₀ '
    'after one year. At the 5th-percentile scenario, puts are typically out-of-the-money — '
    'no payoff. But the investor already paid the full put premium, which drags on the stock '
    'return. So the 5th-percentile return is worse with insurance than without. '
    'This is not a modelling error — it is the correct economic result when drift is high '
    'relative to put cost.'
)

qa_block(
    'Why use the Svensson yield curve rather than a constant rate?',
    'It is more accurate. The Svensson model is the ECB\'s official term-structure model '
    'and captures the slope and curvature of the yield curve. For a certificate maturing '
    'in June 2027, the relevant discount rate changes every day as T_rem shrinks. '
    'A constant rate would introduce a systematic bias. The professor\'s rubric specifically '
    'endorses Svensson for higher marks.'
)

qa_block(
    'What is the CRR binomial tree and why N = 200?',
    'Cox-Ross-Rubinstein (1979): discretize time into N steps, stock moves up by '
    'u = e^{σ√Δt} or down by d = 1/u. Risk-neutral probability p = (e^{(r−q)Δt}−d)/(u−d). '
    'Roll back expected payoffs discounted at r. N = 200 means the barrier is checked at 200 '
    'time points, approximating continuous monitoring. Higher N converges to the true '
    'continuous-barrier price but is slower (~1 second per call).'
)

qa_block(
    'What is EWMA volatility and why λ = 0.94?',
    'Exponentially Weighted Moving Average (RiskMetrics, J.P. Morgan 1994). '
    'Variance updated as: σ²_t = λ·σ²_{t-1} + (1−λ)·r²_t. '
    'λ = 0.94 is the standard RiskMetrics parameter for daily returns. '
    'Each return gets weight (1−0.94) = 6%, decaying exponentially. '
    'It adapts to volatility clusters faster than a simple rolling window.'
)

qa_block(
    'What is Geometric Brownian Motion and what are its limitations?',
    'GBM models the stock as dS = μ·S·dt + σ·S·dW. The exact solution gives log-normally '
    'distributed prices: S_T = S₀·exp[(μ−½σ²)T + σ·W_T]. '
    'Limitations: (1) constant σ — Rheinmetall\'s actual vol ranged 20–70%; '
    '(2) no jumps — defence stocks have discrete geopolitical shocks; '
    '(3) log-normal tails are thinner than empirical fat tails (no kurtosis).'
)

qa_block(
    'Why is Theta POSITIVE for the bonus certificate?',
    'For a vanilla long option, Theta is negative (time decay erodes option value). '
    'For the bonus certificate, the opposite holds in the barrier-sensitive region. '
    'As time passes, the probability of the barrier being breached in the remaining life '
    'falls — the protection becomes more likely to survive intact. '
    'This increases the value of the embedded DO-put, making Theta positive. '
    'It is a structural feature of barrier products, not a mistake.'
)

qa_block(
    'Why is Vega NEGATIVE for the bonus certificate?',
    'Higher volatility increases the probability that Rheinmetall will touch the '
    'knock-in barrier B = 1050 at some point before maturity. If the barrier is '
    'touched, the down-and-out put is knocked out — worth zero. '
    'So higher vol destroys the DO-put\'s value, reducing the total certificate value. '
    'This is the opposite of a long vanilla option, where higher vol always helps.'
)

qa_block(
    'What is the replicating portfolio and why is the equity fraction sometimes above 100%?',
    'The replicating portfolio holds Δ shares and a bond position B = V − Δ·S. '
    'The equity fraction is Δ·S/V. When S approaches the barrier B = 1050, '
    'Delta rises sharply while V is compressed (barrier risk priced in). '
    'So Δ·S/V > 1 — the bond position is negative (borrowed), creating implicit leverage. '
    'This is correct economic behaviour for a barrier product near the barrier.'
)

qa_block(
    'Why is brentq the right solver for Task XI?',
    'brentq (Brent\'s method) finds roots of f(α) = VaR(α) + 0.15 = 0. '
    'It is ideal because: (1) it is guaranteed to converge if f has opposite signs '
    'at the bracket endpoints — confirmed by the monotone scan over 80 α values; '
    '(2) it does not require derivatives — VaR is a simulation quantity, '
    'not analytically differentiable.'
)

qa_block(
    'Your delta time-series plot is missing — what happened?',
    'The file replicating_portfolio_time.png was accidentally deleted before submission '
    '(git history confirms it existed). The plot would show Delta starting near 0.8–0.9 '
    'in September 2025 (RHM at ~1700, well above barrier), rising toward 1.0 as RHM '
    'fell toward the barrier region in spring 2026, with equity fraction peaking above '
    '100% near B = 1050. This is a legitimate deduction.'
)

qa_block(
    'What makes DU2076 a "high complexity" product?',
    'It has a continuous path-dependent barrier: the knock-in feature monitors the stock '
    'every day for ~1.8 years. The payoff switches structure if the barrier is breached. '
    'Decomposition requires a down-and-out put (exotic barrier option) plus a capped forward. '
    'Delta and Vega have non-standard shapes that differ qualitatively from vanilla options. '
    'A discount certificate (underlying minus short call) would be minimum complexity; '
    'a bonus/barrier certificate is explicitly the next tier up per the grading rubric.'
)

qa_block(
    'What is Black-Scholes and when do you use it in the notebook?',
    'Black-Scholes (1973) gives a closed-form European option price: '
    'C = S·e^{−qT}·N(d₁) − K·e^{−rT}·N(d₂), '
    'where d₁ = [ln(S/K) + (r−q+½σ²)T] / (σ√T) and d₂ = d₁ − σ√T. '
    'It is used in: (1) Task II to derive the participation rate α; '
    '(2) Part 3 to price the fictitious put options for portfolio insurance. '
    'For the DU2076 certificate in Part 2, the binomial tree is used because '
    'the barrier condition makes Black-Scholes unsuitable (no closed-form for DO-put with discrete barrier).'
)

qa_block(
    'How did you validate the binomial tree pricing?',
    'The N=200 tree root (1558.89 EUR) on the first trading day exactly matches the '
    'daily loop value after the Bug B2 fix (which corrected the risk-free rate used '
    'in the CSV export). The implied vol back-solver also serves as indirect validation: '
    'when we feed the market-implied vol into the tree, the model recovers the market '
    'price with MAPE = 0.64%, confirming the tree structure is correct.'
)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — BUGS FIXED
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('Bugs Fixed Before Final Submission')

table6 = doc.add_table(rows=1, cols=3)
table6.style = 'Table Grid'
h6 = table6.rows[0].cells
h6[0].text = 'Bug'; h6[1].text = 'What was wrong'; h6[2].text = 'Fix'
for c in h6:
    for r in c.paragraphs[0].runs: r.bold = True
for rd in [
    ('B1 (Cell 56)',
     'GBM applied Itô correction twice: (μ−½σ²)dt where μ already is the log-price drift',
     'Changed to μ·dt + σ√dt·Z — single correction only'),
    ('B2 (Cell 38)',
     'N=200 tree CSV export used fixed RISK_FREE (3.07%) instead of daily Svensson r(T_rem) (1.92%) → root differed from loop by EUR 15',
     'Replaced with get_risk_free_rate(snap_date, T_f)'),
    ('B3 (Cell 60)',
     'Markdown stated μ = 58.57% p.a.; code output was 54.05%',
     'Updated markdown to 54.05% and updated all dependent stats'),
    ('B4 (Cells 66, 70, 74)',
     'Task IX/X/XI discussion cells had stale figures from pre-Bug1 run',
     'Rewritten from actual cell outputs after re-execution'),
]:
    rw = table6.add_row().cells
    for i, v in enumerate(rd): rw[i].text = v

# ── Save ──────────────────────────────────────────────────────────────────────
out = 'B401_Presentation_Notes.docx'
doc.save(out)
print(f'Saved: {out}')
