"""
Generates B401_Assignment_Explanation.docx — a comprehensive grader-prep document
covering theory, code mechanics, and likely Q&A for all 11 tasks.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────

def heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def body(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x8A)
    return p

def qa(question, answer):
    p = doc.add_paragraph()
    r = p.add_run("Q: " + question)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

    p2 = doc.add_paragraph()
    r2 = p2.add_run("A: " + answer)
    r2.font.size = Pt(11)
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.paragraph_format.space_after = Pt(6)

def divider():
    doc.add_paragraph("─" * 80)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("B401 – Continuous-Time Derivatives Pricing")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Take-Home Assignment – Summer Term 2026\nEberhard Karls Universität Tübingen")
r2.font.size = Pt(13)

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("GRADER PREPARATION DOCUMENT\nCode Explanation & Theory Reference")
r3.bold = True
r3.font.size = Pt(15)
r3.font.color.rgb = RGBColor(0x00, 0x4F, 0x9F)

doc.add_paragraph()
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run("Daanish Muzaffar  |  Student ID: 7259472\nDate: 25 June 2026")
r4.font.size = Pt(12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════

heading("Overview & Setup", level=1, color=(0, 79, 159))

body(
    "This document explains the theory, methodology, and code decisions behind every task "
    "in the B401 assignment. It is structured to prepare you for any question a grader might ask — "
    "from high-level conceptual questions to line-by-line code queries.",
    size=11
)

heading("Underlying: Rheinmetall AG (RHM.DE)", level=2)

body(
    "Rheinmetall AG (RHM.DE) was chosen as the underlying for all three parts. It trades on XETRA "
    "and is a leading European defence manufacturer. The same underlying links Part 1 (certificate design) "
    "to Part 2 (valuation of the DZ BANK DU2076 Bonus Cap Certificate on RHM.DE) and Part 3 (portfolio insurance). "
    "This continuity is intentional: it allows a single, internally consistent data pipeline and makes all "
    "valuation results comparable."
)

heading("Data & Global Parameters", level=2)

body("The setup cell loads three groups of data:")
bullet("Daily RHM.DE close prices from Yahoo Finance, 02 Jan 2023 – 01 Jun 2026 (~840 trading days). "
       "The date range was chosen to capture Rheinmetall's full vol regime — the calm 2023-2024 uptrend "
       "and the late-2025 crash — so the full-sample vol estimate is representative.")
bullet("Daily log-returns r_t = ln(S_t / S_{t-1}) from that price series.")
bullet("ECB Svensson yield-curve parameters (daily, Sep 2025 onwards) from the ECB Data Portal.")

body("Key global constants:")
code_block("SIGMA_HIST  = log_returns.std() * sqrt(252)   # full-sample annualised vol (~39.9%)")
code_block("RISK_FREE   = svensson(10.0, **SV_P1)         # 10y ECB Svensson rate (Part 1 only)")
code_block("DIV_YIELD   = 11.50 / S0_CPN                  # Rheinmetall FY2025 dividend yield")
code_block("MATURITY_DATE = 2027-06-18                    # DU2076 maturity date")

qa("Why use the full-sample historical standard deviation for volatility?",
   "It is the closest available proxy to market-implied vol for this product. The full-sample estimate "
   "(~39.9%) sits close to the mean implied vol back-solved from DU2076 market prices (~37.8%), minimising "
   "the structural pricing gap. It also matches the assignment requirement in Task IX: 'prices based on the "
   "historical volatility of the underlying'. Using the same estimator across Parts 2 and 3 ensures "
   "internal consistency.")

qa("Why the Svensson (1994) model for the risk-free rate?",
   "The ECB publishes daily fitted Nelson-Siegel-Svensson parameters for the AAA euro area sovereign curve. "
   "Using this daily, maturity-matched rate (r(T_rem) on each valuation day) is more accurate than a "
   "single fixed benchmark. The formula is: r(t) = b0 + b1*f1 + b2*f2 + b3*f3, where f1, f2, f3 are "
   "functions of the maturity and decay parameters tau1, tau2. For Part 1 a static snapshot from "
   "01 Jun 2026 is used; for Part 2 the rate is re-evaluated daily.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 1
# ══════════════════════════════════════════════════════════════════════════════

heading("PART 1 – Designing Certificates", level=1, color=(0x8B, 0, 0))

# ── Task I ────────────────────────────────────────────────────────────────────
heading("Task I – Investor Profile", level=2)

body(
    "The certificate was designed for German retail investors aged 40–60, 3–5 years from retirement, "
    "who have a structural conviction on the European defence re-armament theme (Germany's €100bn "
    "Sondervermögen, NATO 2% GDP commitments, Rheinmetall's dominant position in armoured vehicles "
    "and munitions) but cannot accept Rheinmetall's ~40% annualised volatility via a direct stock position."
)
body("Key profile characteristics:")
bullet("Age / horizon: 40–60, aligns with the 3-year product tenor")
bullet("Risk tolerance: Low-to-moderate; nominal capital protection is non-negotiable")
bullet("Orientation: Growth-oriented (defence re-rating thesis), constrained by capital preservation")
bullet("Market expectation: Structurally bullish, targeting 15–30% net gain over 3 years")

qa("Why is this investor better served by the certificate than direct stock?",
   "At 40% annualised vol, a direct RHM.DE position can realistically lose 30–50% in a single year — "
   "incompatible with pre-retirement capital preservation. The certificate guarantees nominal capital "
   "return at maturity (the ZCB leg), provides partial upside participation (the ATM call spread), and "
   "eliminates the catastrophic loss scenario, at the cost of capping the maximum gain at 30% and "
   "foregoing dividends.")

qa("Why not a bond or money-market fund?",
   "The investor has a specific conviction on Rheinmetall. A risk-free bond does not express that view "
   "at all. The certificate gives genuine participation in the upside — 83.8% of gains up to 30% — "
   "which a bond cannot provide, while retaining the same nominal floor.")

# ── Task II ───────────────────────────────────────────────────────────────────
heading("Task II – Product Design: Capped Capital Protected Participation Note", level=2)

body(
    "The product is a three-component structured certificate. The key design insight is to use high "
    "Rheinmetall volatility to the investor's advantage: elevated vol makes both the long and short call "
    "expensive, but the spread (long ATM minus short OTM) is less vol-sensitive than either individual "
    "option, producing a higher participation rate than vol alone might suggest."
)

heading("Components", level=3)
bullet("Zero-Coupon Bond (ZCB): face value = S₀, maturity T = 3 years. "
       "Cost = S₀ × e^(−r×T). Guarantees full nominal return at expiry.")
bullet("Long ATM European Call: K = S₀ (at-the-money). Provides upside participation.")
bullet("Short OTM European Call: K = 1.30×S₀ (cap level). Caps gains at 30%; premium received "
       "funds a higher participation rate α than the option budget alone would allow.")

heading("At-Par Design Condition", level=3)

body("The product must issue at par, so:")
code_block("S₀  =  ZCB₀  +  α × (Call_ATM − Call_Cap)")
code_block("α   =  (S₀ − ZCB₀) / (Call_ATM − Call_Cap)")

body("With r(3y) ≈ 1.78% (Svensson), σ = 39.9%, S₀ ≈ €1,207:")
bullet("ZCB₀ = S₀ × e^(−1.78% × 3) ≈ €1,115.79")
bullet("ATM call (K=1,207) ≈ €346.90")
bullet("OTM call (K=1,569) ≈ €238.05")
bullet("Spread = €108.86")
bullet("Participation rate α = (1,207 − 1,115.79) / 108.86 ≈ 83.8%")

qa("How is Black-Scholes applied to compute α?",
   "The standard GBM-based closed-form formula with continuous dividend yield q: "
   "C = S×e^(−qT)×N(d1) − K×e^(−rT)×N(d2), where "
   "d1 = [ln(S/K) + (r−q+σ²/2)×T] / (σ√T), d2 = d1 − σ√T. "
   "Both the ATM and OTM calls are priced this way; the difference is their moneyness. "
   "The dividend yield q = 0.95% is the Rheinmetall FY2025 yield (€11.50 / S₀).")

qa("Why cap at 130% and not 120% or 150%?",
   "At σ = 39.9%, a 130% strike is moderately OTM and commands meaningful premium (≈€238). "
   "Going lower (120%) brings in more premium but caps the investor too tightly — a 15% scenario "
   "would already approach the cap. Going higher (150%) generates less premium, producing a lower α. "
   "At 130% the participation rate (83.8%) and max gain (≈25%) align with the investor's 15–30% target "
   "scenario while the ZCB still fits the budget.")

qa("What happens if rates rise significantly after issuance?",
   "The ZCB is the most rate-sensitive component: its present value = S₀×e^(−rT) falls when rates rise. "
   "At a fixed issue price of S₀, a higher r would mean more option budget and thus a higher α — but "
   "the product is already issued at fixed terms. Secondary market value falls with rates (rising rates "
   "compress ZCB prices), which would hurt mark-to-market value. The certificate does not mark-to-market "
   "for a retail buy-and-hold investor, so this only matters if sold before maturity.")

# ── Task III ──────────────────────────────────────────────────────────────────
heading("Task III – Payoff Profile", level=2)

body("The payoff function at maturity (in EUR, as a function of terminal price S_T):")
code_block("Payoff(S_T) = S₀                            if S_T < K    (protection floor)")
code_block("           = S₀ + α × (S_T − K)             if K ≤ S_T ≤ Cap  (participation zone)")
code_block("           = S₀ + α × (Cap − K)             if S_T > Cap  (maximum payoff, capped)")

body("The graph shows three clearly labelled regimes:")
bullet("BEARISH (S_T < S₀): CPN pays S₀, outperforming stock and achieving the protection objective.")
bullet("MODERATELY BULLISH (S₀ < S_T ≤ Cap): CPN participates at 83.8%, stock slightly ahead. "
       "Both lines slope upward; the gap between them is the cost of protection.")
bullet("STRONGLY BULLISH (S_T > Cap): CPN is capped at max payoff ≈ €1,307; stock outperforms. "
       "This is the trade-off — the investor gave up extreme upside for the downside floor.")

qa("Where does the CPN break even against the stock?",
   "At S_T = S₀ (the protection floor) the CPN returns S₀ while a direct stock holder also has S₀ "
   "(ignoring dividends). Below S₀, the CPN outperforms. Above S₀ but in the participation zone, "
   "the CPN slightly underperforms (the 16.2% missing participation is the insurance cost). "
   "There is also a break-even vs. the risk-free bond at S_T ≈ bond_mature/α + K, above which "
   "the certificate beats the bond in upside terms.")

qa("Does the certificate pay dividends?",
   "No. The ZCB and call structure replicate equity payoff but not dividend cash flows. "
   "The dividend yield q enters the Black-Scholes formula to adjust the forward price but "
   "the investor does not receive Rheinmetall's €11.50 FY2025 dividend. This is one reason "
   "the certificate can afford to offer the 100% capital protection — the issuer keeps dividends "
   "to partly fund the hedge.")

# ── Task IV ───────────────────────────────────────────────────────────────────
heading("Task IV – Market Size Estimation", level=2)

body("The market size is estimated from the DDV (Deutscher Derivate Verband) 2023 annual statistics:")
bullet("Total German structured product market: EUR 112 bn")
bullet("Kapitalschutz-Zertifikate segment (capital protected): 3.8% → EUR 4.25 bn")
bullet("Participation notes sub-segment (~55% of Kapitalschutz): EUR 2.34 bn")
bullet("~130 active issues (≈6 major issuers × 20 products each) → average issue ≈ EUR 33 mn")

qa("Why is the Kapitalschutz segment relatively small (3.8%)?",
   "Capital protection requires a large ZCB allocation, which is expensive in higher-rate environments. "
   "Discount certificates (18%) and Index/Partizipation (34%) are cheaper to construct because they "
   "don't guarantee nominal return. The Kapitalschutz segment has grown since 2020 as the post-COVID "
   "rate environment and equity participation demand increased, but it remains niche because the "
   "required return from the investor's perspective is low (capped upside, no dividends).")

qa("Who are the main issuers?",
   "DZ BANK, Commerzbank, LBBW, Goldman Sachs, HypoVereinsbank/UniCredit, Société Générale. "
   "These are MiFID II-registered structured product issuers with distribution networks covering "
   "German retail banks. DZ BANK is also the issuer of DU2076 in Part 2.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 2
# ══════════════════════════════════════════════════════════════════════════════

heading("PART 2 – Valuation of Certificates", level=1, color=(0x8B, 0, 0))

heading("Product: DU2076 Bonus Cap Certificate (DZ BANK)", level=2)

body(
    "DU2076 (ISIN DE000DU20767) is a Bonus Cap Certificate issued by DZ BANK AG on Rheinmetall AG (RHM.DE). "
    "It was listed on 4 September 2025 and matures on 18 June 2027."
)
body("Key terms from the official DZ BANK KID / product page:")
bullet("Barrier B = €1,050 (continuous, observed daily during the product's life)")
bullet("Bonus / Cap level K = €2,000")
bullet("Participation ratio: 1:1 (Bezugsverhältnis = 1.00)")
bullet("Barrier NOT breached as of 01 Jun 2026 (distance to barrier: ~12.9% from €1,207)")

body("Payoff at maturity:")
code_block("V_T = 2,000  if S_t > 1,050 for ALL t in [04 Sep 2025, 18 Jun 2027]")
code_block("V_T = min(S_T, 2,000)  if barrier is breached at any point")

qa("Why is DU2076 a good product for Part 2?",
   "It uses the same underlying as Part 1 (Rheinmetall), has a well-documented public data source "
   "(DZ BANK issuer page, EUWAX, onvista.de), and contains a path-dependent feature (continuous barrier) "
   "that motivates the binomial tree over Black-Scholes, making the technical valuation more interesting. "
   "It also has an active secondary market with 183 days of observable market prices for comparison.")

# ── Pricing Engine ────────────────────────────────────────────────────────────
heading("Pricing Engine: Binomial Tree", level=2)

heading("Why Not Black-Scholes?", level=3)

body(
    "Black-Scholes prices options whose payoff depends ONLY on the terminal price S_T. "
    "It models the stock jumping directly from S_0 to S_T (via the log-normal terminal distribution) "
    "and has no representation of intermediate prices. The DU2076 barrier condition is path-dependent: "
    "the bonus is lost if S_t ≤ 1,050 at ANY point during the observation window — a statement about "
    "the entire price path, not the endpoint. Black-Scholes cannot evaluate this condition."
)

heading("Binomial Tree Mechanics", level=3)

body("The CRR (Cox-Ross-Rubinstein) binomial tree discretises time into N = 200 steps of size Δt = T/N.")
body("At each node the stock moves up or down:")
code_block("u = exp(σ × sqrt(Δt))    (up factor)")
code_block("d = 1/u                   (down factor; tree is recombining since u×d = 1)")
code_block("q = (exp((r−θ)×Δt) − d) / (u − d)   (risk-neutral up-probability)")

body("where r is the Svensson rate, θ = DIV_YIELD is the continuous dividend yield.")

body("The certificate is decomposed into two components valued in a single backward-induction pass:")
bullet("Down-and-out put (DO-put): pays max(K−S_T, 0) at maturity IF the stock never touched "
       "the barrier B. At every backward step, nodes where S_{n,j} ≤ B are set to zero (absorbed). "
       "This correctly discounts paths that knocked out.")
bullet("Capped forward: pays min(S_T, K_cap) at maturity with NO barrier condition. "
       "Equivalent to holding the stock and selling a call at K_cap.")
bullet("Certificate price = DO-put + Capped forward")

body("Backward induction step (from terminal to today):")
code_block("V[n, j] = disc_factor × (q × V[n+1, j+1] + (1−q) × V[n+1, j])")
code_block("  then: V_do[n, j] = 0  if S[n, j] ≤ B  (barrier absorption)")

qa("Why N = 200 steps?",
   "Convergence analysis for barrier options in binomial trees typically requires N ≥ 100 to keep "
   "the pricing error below 1 EUR. N = 200 provides reliable convergence while keeping runtime "
   "under 1 second per pricing call with NumPy vectorisation. For the Greeks and the Task XI scan, "
   "a faster N = 50 approximation is used; for the main valuation and binomial tree export, N = 200.")

qa("Why is the tree recombining (u×d = 1)?",
   "A recombining tree has N+1 terminal nodes instead of 2^N. With N=200, that is 201 nodes instead "
   "of 2^200 — computationally feasible. The CRR choice u = exp(σ√Δt), d = 1/u guarantees recombination "
   "because one up-move followed by one down-move returns to the same price: S×u×d = S×1 = S.")

qa("How is the barrier implemented?",
   "At every backward-induction step n, after computing the rolled-back DO-put value for each node, "
   "we check S[n,j] ≤ B and force V_do[n,j] = 0. This absorbs knocked-out branches permanently — "
   "once a branch crosses the barrier, it contributes zero to the backward sum for all earlier steps. "
   "The capped forward is NOT affected by the barrier (it survives regardless).")

# ── Task V ────────────────────────────────────────────────────────────────────
heading("Task V – Daily Valuation (183 days, Sep 2025 – Jun 2026)", level=2)

body("For each of the 183 trading days in the window:")
bullet("T_rem = (MATURITY_DATE − date).days / 365  — remaining time to maturity in years")
bullet("r_today = get_risk_free_rate(date, T_rem)  — daily maturity-matched Svensson rate")
bullet("sigma_est = SIGMA_HIST = 39.9%  — constant (full-sample historical std dev)")
bullet("Model price = bonus_cap_cert_price(S, K=2000, B=1050, r, q, sigma, T_rem)")

body("Market prices are loaded from data/cert_prices.csv (onvista.de historical export of DU2076 "
     "closing prices, semicolon-separated, German decimal format).")

body("Error metrics (model vs market):")
code_block("ME   = mean(model − market)  = −43.64 EUR")
code_block("MAE  = mean(|model − market|) = 46.67 EUR")
code_block("RMSE = sqrt(mean((model − market)²)) = 57.51 EUR")
code_block("MAPE = mean(|error| / market) × 100   = 2.90%")

body("A second model is computed by back-solving the market-implied volatility each day (Brent's "
     "root-finding on the N=50 tree) and re-pricing at N=200. This IV model achieves ME ≈ −8.89 EUR, "
     "MAPE ≈ 0.64% — confirming the hist-vol gap is mainly a vol-level difference, not a model error.")

qa("Why is the pricing error negative (model < market) in early months?",
   "In September 2025 Rheinmetall was near €1,700–1,800 — far above the €1,050 barrier — and the "
   "market's implied vol was higher than the 39.9% historical estimate (the market was pricing "
   "significant jump risk from the then-active vol regime). Higher implied vol pushes the DO-put up "
   "(barrier is more likely to be approached under higher vol), raising the market price above the "
   "historical-vol model. The November 2025 crash at ~€1,570 was the worst day for this discrepancy.")

qa("Why does the error flip to positive (model > market) in late May 2026?",
   "By May 2026 Rheinmetall had fallen to ~€1,120–1,240, close to the barrier. Near the barrier, "
   "the DO-put's value is very sensitive to the barrier probability. The market may be pricing "
   "a slightly lower vol (barrier already close, put is nearly ATM) while the model uses the "
   "constant historical vol. Also, issuer bid-ask spreads widen near barrier proximity, making "
   "observed market prices slightly lower than fair value.")

qa("Why use a constant historical vol rather than a rolling window?",
   "The assignment asks for a single consistent vol estimate. A rolling window would produce different "
   "vol estimates each day, complicating comparison. The full-sample estimate is also the closest to "
   "what an investor would observe at purchase date. The IV back-solve (second model) confirms that "
   "using the correct implied vol per day would reduce MAPE to 0.64%.")

# ── Task VI ───────────────────────────────────────────────────────────────────
heading("Task VI – Greeks / Sensitivity Analysis", level=2)

body("Since DU2076 has no closed-form Greeks, all sensitivities are computed numerically via "
     "central finite differences on the N=200 binomial tree:")
code_block("Δ (Delta)  = [V(S+ε) − V(S−ε)] / (2ε×S)          ε = 1% × S₀")
code_block("Γ (Gamma)  = [V(S+ε) − 2V(S) + V(S−ε)] / (ε×S)²  same ε")
code_block("V (Vega)   = [V(σ+h) − V(σ−h)] / (2h) / 100       h = 2% vol")
code_block("Θ (Theta)  = [V(T−Δt) − V(T)] / Δt / 365          Δt = 1/252")

body("The cross-section plots these Greeks across S from 0.55×B to 1.55×K, at fixed "
     "T_rem ≈ 1.78y (first DU2076 trading day) and σ = 39.9%.")

heading("Interpretation of Each Greek", level=3)
bullet("Delta: At high S (far above B), Δ ≈ 0.2–0.3 — the certificate moves only ~25 cents per €1 "
       "move in Rheinmetall (most value comes from the guaranteed bonus, not the stock leg). "
       "Near the barrier, Δ spikes toward 1.0 — the certificate tracks the stock almost 1:1 because "
       "whether the barrier is crossed becomes the dominant value driver.")
bullet("Gamma: Negative near S_ref (approaching cap) because the certificate's capped structure "
       "means Delta falls as S rises above K=2,000 (like being short an OTM call). Near the barrier, "
       "Gamma spikes sharply negative: Delta is changing rapidly as S approaches B, making "
       "the position hard to hedge — the issuer faces serious gamma risk near B=1,050.")
bullet("Vega: Negative (≈ −1.18 EUR per 1% vol increase). Two simultaneous effects: (1) "
       "the capped forward is effectively short a call at K=2,000; higher vol raises that short "
       "call's value, hurting the position. (2) Higher vol increases barrier-breach probability, "
       "reducing the expected value of the DO-put bonus. Both push in the same direction.")
bullet("Theta: Positive (≈ +0.31 EUR/day at S_ref). Deep in the bonus zone, each passing day "
       "without a barrier breach makes the €2,000 bonus payout more certain. The certificate gains "
       "value from time decay — opposite to a vanilla option. Think of it like a deeply ITM European put.")

qa("Why use a 2% vol bump for Vega instead of the standard 1%?",
   "The binomial tree introduces discrete grid noise because lattice nodes shift as sigma changes. "
   "With a 1% bump the noise can exceed the true Vega signal, producing unreliable results. A 2% bump "
   "is wide enough to average over the grid noise while still being small enough to be a valid first-order "
   "approximation. This is a known issue with binomial-tree Greeks and is documented in Hull "
   "(Options, Futures, and Other Derivatives, 11th ed., Chapter 19).")

qa("If Vega is negative, should the issuer hedge by selling vol?",
   "The issuer is short the certificate (they sold it). Being short a negative-Vega position means "
   "the issuer is long Vega — they benefit from higher vol. Their hedge is therefore to sell vol "
   "(e.g., sell variance swaps or short straddles) to neutralise their net Vega exposure. "
   "This is the standard dealer hedging logic for reverse-Vega products.")

# ── Task VII ──────────────────────────────────────────────────────────────────
heading("Task VII – Replicating Portfolio", level=2)

body("The replicating portfolio at each date t holds Δ(t) shares and a bond position B(t):")
code_block("V(t)    = Δ(t) × S(t) + B(t)")
code_block("B(t)    = V(t) − Δ(t) × S(t)       (bond leg, can be negative)")
code_block("Equity fraction = Δ(t) × S(t) / V(t)")

body("Δ(t) is computed daily from the binomial tree central-difference formula (same as Delta in Task VI). "
     "V(t) is the model price on that day. Both use constant σ = 39.9% and the daily Svensson rate.")

heading("Time Series (Sep 2025 – Jun 2026)", level=3)
bullet("Delta started at ~0.25 in Sep 2025 (Rheinmetall far above barrier at €1,708). "
       "The certificate behaved like a mildly leveraged capped forward.")
bullet("As Rheinmetall fell through 2026, Delta drifted toward 0.9–1.0 near the barrier region. "
       "Near B=1,050, the certificate was essentially a 1:1 stock position.")
bullet("The equity fraction followed the same pattern: started ~25–30%, rose to 80–95% near "
       "the barrier approach. Bond leg shrank as stock sensitivity increased.")

heading("Static Cross-Section (Equity Fraction vs S)", level=3)
bullet("Below B=1,050: certificate has knocked out, reverts to capped forward. "
       "Equity fraction rises steeply as S falls — full downside exposure.")
bullet("Near B: fraction peaks sharply (Delta → 1). Hardest hedging region for the issuer.")
bullet("In bonus zone (B < S < K): fraction is moderate (20–60%). Main structural zone.")
bullet("Near/above K=2,000: fraction falls as certificate value is increasingly bond-like "
       "(guaranteed bonus nearly certain, little incremental stock sensitivity).")

qa("Why does equity fraction exceed 100% near the barrier?",
   "When Delta ≈ 1.0 and the certificate value V is falling (because barrier breach is imminent), "
   "the ratio Δ×S/V can exceed 1. In portfolio terms: the replicating portfolio is short the bond "
   "(B(t) < 0) and long more than one unit of stock. This is the 'leveraged near-barrier' regime. "
   "It is an accurate representation — the issuer must hold more than 1 share per certificate to "
   "replicate the payoff because at barrier breach the payoff discontinuously shifts from €2,000 to S_T.")

qa("What is the key difference between the CPN (Part 1) and DU2076 (Part 2) in falling markets?",
   "The CPN (capital protected) has Delta → 0 as S falls below K=S₀, reducing stock exposure. "
   "DU2076 has Delta → 1 as S approaches B=1,050, INCREASING stock exposure. This is the "
   "structural reversal: protection notes become more bond-like in downturns, while barrier "
   "certificates become more stock-like at exactly the worst time. An investor holding DU2076 "
   "in a sustained decline effectively holds a leveraged stock position near B — the opposite of "
   "what 'protection' might imply.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 3
# ══════════════════════════════════════════════════════════════════════════════

heading("PART 3 – Portfolio Insurance Strategy", level=1, color=(0x8B, 0, 0))

body("Setup: W₀ = €10,000, T* = 1 year, underlying = Rheinmetall AG (RHM.DE). "
     "Put options are 'fictitious' — priced via Black-Scholes at issuance (t=0) using historical vol.")

heading("Monte Carlo Simulation – GBM", level=2)

body("50,000 independent 1-year price paths are simulated at daily frequency (252 steps) "
     "under Geometric Brownian Motion:")
code_block("S_{t+Δt} = S_t × exp[(μ − σ²/2)×Δt + σ×sqrt(Δt)×Z_t]    Z_t ~ N(0,1)")
body("Parameters:")
code_block(f"S₀    = current RHM.DE spot (€1,207)      μ = 54.05% p.a. (historical mean log-return × 252)")
code_block(f"σ     = 39.88% (SIGMA_HIST)               r = 3.07% (ECB Svensson 10y)")
code_block(f"Seed  = 2026 (for reproducibility)        N_STEPS = 252, N_PATHS = 50,000")

body("The same path set is shared by Tasks VIII–XI so differences in outcomes are purely "
     "strategy-driven, not Monte Carlo noise.")

qa("Why use the historical drift μ = 54% and not the risk-neutral drift r = 3%?",
   "Performance analysis (Tasks VIII–XI) evaluates the REAL-WORLD probability distribution of "
   "returns — what the investor actually experiences, not the no-arbitrage risk-neutral measure. "
   "Under P-measure we simulate with the estimated physical drift μ. The risk-neutral measure "
   "(r = 3%) is used only for PRICING the options at t=0. This is standard: price under Q, simulate "
   "under P.")

qa("What are the limitations of GBM for Rheinmetall?",
   "Three main limitations: (1) Constant volatility — RHM.DE's realised vol ranged from <20% to >70% "
   "over 2023–2026; a stochastic vol model (Heston, SABR) would be more accurate. (2) Continuous "
   "paths — GBM cannot produce the late-Nov 2025 ~16% single-session crash; jump-diffusion (Merton 1976) "
   "is the natural extension. (3) Historical drift — 54% is specific to the 2022-2026 defence re-rating "
   "regime; it is not a realistic long-run forecast (a sensible equity risk premium would be 6–9%).")

# ── Task VIII ─────────────────────────────────────────────────────────────────
heading("Task VIII – Performance Analysis (No Risk Management)", level=2)

body("The unhedged position invests 100% of W₀ = €10,000 in Rheinmetall stock. "
     "Returns are log-returns of terminal wealth: r = ln(S_T / S₀).")

body("Key statistics from the 50,000-path simulation:")
bullet("Mean return: ~54% (equals drift μ, as expected for GBM)")
bullet("Standard deviation: ~39.9% (equals σ by construction)")
bullet("Skewness: ~0.007 (near-zero — GBM log-returns are symmetric)")
bullet("95% VaR: ~11.3% loss at the 5th percentile")
bullet("95% CVaR: ~20–22% loss (average of worst 5% outcomes)")
bullet("P(loss): ~8.7% of paths end below starting price")

qa("Why is the return distribution right-skewed for the stock price but symmetric for log-returns?",
   "GBM generates log-normal stock PRICES (right-skewed, non-negative) but normal log-RETURNS "
   "(symmetric by construction). The histogram plots log-returns, so it appears approximately "
   "symmetric. The slight positive skewness in price space comes from the multiplicative nature "
   "of GBM: up moves compound more than down moves of equal magnitude.")

qa("What does 95% VaR = 11.3% mean in practice?",
   "With W₀ = €10,000 and a 1-year horizon, there is a 5% probability that the portfolio loses "
   "more than 11.3% of its value, i.e., more than €1,130. Equivalently, the 5th-percentile "
   "terminal wealth is ≈ €8,870. This is the baseline risk measure that Tasks IX–XI attempt to "
   "manage via put options.")

# ── Task IX ───────────────────────────────────────────────────────────────────
heading("Task IX – Portfolio Insurance with Put Options", level=2)

body("A fraction α of W₀ is invested in European put options at issuance, (1−α) in stock:")
code_block("n_puts   = (α × W₀) / put_price(S₀, K, r, q, σ, T)")
code_block("n_shares = ((1−α) × W₀) / S₀")
code_block("Wealth_T = n_puts × max(K − S_T, 0) + n_shares × S_T")
code_block("Return   = ln(Wealth_T / W₀)")

body("Put price = Black-Scholes European put at issuance: P = K×e^(−rT)×N(−d2) − S×e^(−qT)×N(−d1)")
body("Grid explored: α ∈ {5%, 10%, 15%, 20%} × K ∈ {90%, 95%, 100%, 105%} of S₀ → 16 combinations.")

heading("Main Finding: Insurance RAISES VaR (under high drift)", level=3)

body(
    "Counter-intuitively, adding puts increases the 95% VaR rather than reducing it for "
    "K = 90–100% at all allocation levels. The mechanism:"
)
bullet("With μ = 54%, only ~8.7% of paths end below S₀. At the 5th percentile, the stock "
       "has usually risen 5–15% — the put expires out-of-the-money.")
bullet("The put premium is paid at t=0 regardless. Higher α → more premium paid → worse "
       "5th-percentile outcome (premium burn with no offsetting payoff).")
bullet("This does NOT mean puts are useless. They genuinely protect in the worst outcomes "
       "(below K), but those outcomes occur at probabilities below 5% in this drift regime, "
       "so they don't show up in the 95% VaR.")

body("Strike reading (columns): Higher-strike puts (K=105%) have better VaR at given α because "
     "in the ~8.7% of falling paths they deliver more payoff. K=105% is nearly ATM, activating "
     "more often.")
body("Allocation reading (rows): VaR worsens monotonically with α for K=90–100%. "
     "The premium cost dominates at the VaR quantile where puts are worthless.")

qa("At what drift would insurance genuinely improve VaR?",
   "At a conventional long-run drift of μ ≈ 6–9% (equity risk premium over r), the stock falls "
   "at the 5th percentile — the put would activate and provide genuine tail protection. "
   "VaR would decrease with α, making alpha* from Task XI a genuine target, not a regulatory ceiling. "
   "The current result is specific to Rheinmetall's extraordinary 2022–2026 return environment.")

qa("Why are 'fictitious' puts used instead of market prices?",
   "Task IX specifies 'fictitious puts whose prices are based on the historical volatility'. "
   "Real DU2076-style certificates don't have liquid European-put markets at arbitrary strikes. "
   "Using BS prices at σ = SIGMA_HIST ensures internal consistency and allows the full α×K grid "
   "to be evaluated at clean, comparable prices. In practice an issuer would use implied vol; "
   "the vol sensitivity is explored in Task X.")

# ── Task X ────────────────────────────────────────────────────────────────────
heading("Task X – Stress Scenario Analysis", level=2)

body("Chosen allocation: α = 10%, K = 95% of S₀ (a mid-range, economically sensible choice). "
     "Four scenarios are tested:")

bullet("Baseline: standard GBM, σ_pricing = SIGMA_HIST ≈ 39.9%.  VaR = 16.00%")
bullet("Vol +5pp: pricing vol raised to σ+0.05 at issuance (FRTB/Basel IV vega stress magnitude).  "
       "VaR = 16.77% — put premium increases, worsening premium drag at VaR quantile.")
bullet("Vol −5pp: pricing vol lowered to σ−0.05.  VaR = 13.62% — cheaper puts, less drag, "
       "only scenario where α=10% is VaR-compliant.")
bullet("−20% price shock at T/2: mid-horizon crash (PRIIPs KID stress scenario methodology).  "
       "VaR = 17.07%, P(loss) triples to 29.8%, mean return drops from 44.7% to 24.8%.")

body("The ±5pp magnitude follows Basel FRTB (BCBS Jan 2019, §21.83); the −20% shock follows "
     "the PRIIPs KID adverse scenario for high-vol equities (EU Regulation 1286/2014).")

qa("What does the price shock scenario tell us about the put?",
   "The put struck at K = 95%×S₀ was initially OTM. After a −20% shock at T/2, the stock is "
   "~20% below S₀ and the put is deep ITM. It provides significant payoff — but the 90% stock "
   "position simultaneously suffers a large loss. Mean return drops to 24.8%, showing the put "
   "helps but doesn't fully offset the crash (because only 10% was allocated to puts). "
   "P(loss) rises from 13.7% to 29.8% because many paths that would have recovered from the crash "
   "in the baseline are now time-constrained (only T/2 of recovery time remains).")

qa("Why does vol +5pp worsen VaR more than the baseline?",
   "Higher vol at issuance makes puts more expensive (Black-Scholes Vega is positive for puts). "
   "At the 5th percentile the put still likely expires worthless (stock up moderately), so the "
   "extra premium paid is pure drag. The vol scenario tests 'what if I overpay for the options' — "
   "the answer is worse VaR without better protection at the quantile that matters.")

# ── Task XI ───────────────────────────────────────────────────────────────────
heading("Task XI – Capital Requirement (VaR Constraint)", level=2)

body("The binding constraint: 95% one-year VaR cannot exceed 15%. "
     "Find α* such that VaR(α*) = exactly −15% (fully utilised constraint).")

body("Method: scipy.optimize.brentq, applied to the function f(α) = VaR(α) − (−0.15). "
     "Brent's method requires a sign change: verified by preliminary scan over α ∈ [0.1%, 40%] "
     "that VaR is monotone increasing (VaR worsens as α rises in this drift regime). "
     "The root α* = 7.90% is where VaR = exactly −15.00%.")

body("At α* = 7.90%, K = 95%×S₀:")
code_block("Put allocation:    €790   (7.90% × €10,000)")
code_block("Stock allocation:  €9,210 (92.10% × €10,000)")
code_block("Expected return:   ~46.73%")
code_block("Sharpe ratio:      ~1.148")
code_block("95% VaR:           15.00% (binding)")

qa("Why is the VaR-minimising allocation α = 0, not α* = 7.90%?",
   "Because VaR increases with α in this drift environment (as shown in Task IX). The unhedged "
   "VaR (11.32%) is already better than any insured portfolio's VaR. α* = 7.90% is not an optimum — "
   "it is the allocation that exactly saturates the regulatory ceiling from below. "
   "If the constraint is 15%, all α ≤ 7.90% are compliant; α* is the maximum compliant allocation.")

qa("What would change at a more conservative drift assumption?",
   "At μ = 8% (a realistic long-run equity risk premium), the 5th percentile of the stock's "
   "terminal distribution would be a loss, puts would activate at the VaR quantile, and VaR "
   "would DECREASE with α. In that scenario α* would be a genuinely optimal protection level, "
   "not a regulatory ceiling. The drift sensitivity table in the notebook makes this explicit.")

qa("Why use Brent's method (brentq) for the root finding?",
   "Brent's method is a robust bracketed root-finding algorithm combining bisection, secant, and "
   "inverse quadratic interpolation. It is guaranteed to converge to a root within the bracket "
   "[a, b] as long as f(a) and f(b) have opposite signs — which we verify by the preliminary "
   "scan. It is faster than pure bisection and more robust than Newton's method (no derivatives "
   "needed, no divergence risk). scipy.optimize.brentq is the standard implementation.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# THEORY PRIMER
# ══════════════════════════════════════════════════════════════════════════════

heading("Theory Primer – Key Formulas at a Glance", level=1, color=(0, 79, 159))

heading("Black-Scholes (European vanilla options)", level=2)
code_block("Call = S×e^(−qT)×N(d1) − K×e^(−rT)×N(d2)")
code_block("Put  = K×e^(−rT)×N(−d2) − S×e^(−qT)×N(−d1)")
code_block("d1   = [ln(S/K) + (r − q + σ²/2)×T] / (σ×√T)")
code_block("d2   = d1 − σ×√T")
body("Parameters: S=spot, K=strike, r=risk-free rate, q=dividend yield, σ=vol, T=time to maturity, "
     "N(·)=standard normal CDF.")

heading("Black-Scholes Greeks (Vanilla Call)", level=2)
code_block("Delta = N(d1)")
code_block("Gamma = N'(d1) / (S × σ × √T)           N'(·) = standard normal PDF")
code_block("Vega  = S × √T × N'(d1)                  (per unit vol change)")
code_block("Theta = −[S×N'(d1)×σ / (2√T)] − r×K×e^(−rT)×N(d2)")

heading("CRR Binomial Tree", level=2)
code_block("u = exp(σ√Δt),   d = 1/u,   Δt = T/N")
code_block("q = [exp((r−θ)Δt) − d] / (u − d)    risk-neutral up-probability")
code_block("disc = exp(−r×Δt)                    one-step discount factor")
code_block("V[n,j] = disc × (q×V[n+1,j+1] + (1−q)×V[n+1,j])")
code_block("Barrier knock-out: V_do[n,j] = 0 if S[n,j] ≤ B")

heading("Svensson (1994) Yield Curve", level=2)
code_block("r(t) = (b0 + b1×f1 + b2×f2 + b3×f3) / 100")
code_block("f1 = [1 − exp(−t/τ1)] / (t/τ1)")
code_block("f2 = f1 − exp(−t/τ1)")
code_block("f3 = [1 − exp(−t/τ2)] / (t/τ2) − exp(−t/τ2)")
body("Parameters b0, b1, b2, b3, τ1, τ2 from ECB Data Portal daily publication.")

heading("GBM Monte Carlo", level=2)
code_block("S_{t+Δt} = S_t × exp[(μ − σ²/2)Δt + σ√Δt × Z]    Z ~ N(0,1)")
code_block("log-return: r = ln(S_T / S₀)  ~  N[(μ−σ²/2)T, σ²T]")
body("μ = physical drift (historical), σ = volatility. For pricing use risk-neutral drift r.")

heading("VaR and CVaR", level=2)
code_block("95% VaR  = −percentile(returns, 5)     (5th percentile loss)")
code_block("95% CVaR = −mean(returns where returns ≤ VaR)   (expected shortfall)")
code_block("Sharpe   = (mean_return − r) / std_return")

heading("Replicating Portfolio", level=2)
code_block("V(t) = Δ(t)×S(t) + B(t)         B(t) = bond leg (can be negative)")
code_block("Δ(t) = [C_u − C_d] / [S_u − S_d]  (binomial tree delta at node closest to S_t)")
code_block("Equity fraction = Δ(t)×S(t) / V(t)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# COMMON GRADER QUESTIONS
# ══════════════════════════════════════════════════════════════════════════════

heading("Likely Grader Questions & Answers (Summary)", level=1, color=(0, 79, 159))

body("These questions may be asked in any order. Answers above go into depth; "
     "these are concise reference versions for quick recall during a presentation.", italic=True)

qa("What is the underlying and why did you choose it?",
   "Rheinmetall AG (RHM.DE, XETRA). Chosen because it is a high-profile, high-vol German large-cap "
   "with a strong current narrative (European rearmament). It has liquid market data, an actively "
   "traded DZ BANK certificate (DU2076) for Part 2 valuation comparison, and justifies both the "
   "capital protection product (high vol, conservative investor) and the portfolio insurance analysis "
   "(40% vol makes insurance decisions non-trivial).")

qa("What is the difference between historical vol and implied vol? Which did you use and why?",
   "Historical vol is computed from past price data: σ_hist = std(log-returns) × √252. "
   "Implied vol is backed out from current market option prices via the BS formula. "
   "I used historical vol (39.9%) as the primary pricing vol for internal consistency with "
   "Task IX which requires 'historical volatility of the underlying'. For comparison, I back-solved "
   "the daily implied vol from DU2076 market prices — the IV model achieves MAPE of 0.64% vs 2.90% "
   "for hist-vol, confirming the hist-vol gap is primarily a vol level difference.")

qa("What is the risk-neutral probability and why do we use it for pricing?",
   "q = [exp((r−θ)Δt) − d] / (u − d). It is not the real probability of an up-move. "
   "It is the probability that makes the discounted expected stock price equal to the current price "
   "(no-arbitrage condition). Under q, all assets earn the risk-free rate — so we can price any "
   "derivative by discounting its expected payoff under q. Using the real probability would give "
   "the wrong price unless we also adjusted the discount rate for risk. Risk-neutral pricing avoids "
   "specifying the investor's risk preference.")

qa("How does the certificate payoff differ in falling vs rising markets?",
   "For the CPN (Part 1): in falling markets (S_T < S₀) the ZCB pays out and the call expires "
   "worthless → investor recovers S₀. In rising markets the call delivers α×(S_T−S₀) capped at α×30%. "
   "For DU2076 (Part 2): in rising markets (no barrier breach) the investor gets €2,000 regardless "
   "of how high the stock goes (capped upside). In falling markets if barrier is NOT breached the "
   "investor still gets €2,000 — a free lunch of sorts. If barrier IS breached, the investor gets "
   "min(S_T, 2,000) — full downside exposure with only an upside cap.")

qa("Why does the equity fraction exceed 100% near the barrier?",
   "Because Δ → 1 as S → B (the certificate mirrors the stock 1:1 near knock-out) while V(t) "
   "is simultaneously falling (the bonus is at risk). The ratio Δ×S/V overshoots 1. The bond leg "
   "B(t) = V(t) − Δ(t)×S(t) turns negative — the replicating portfolio is short a bond (leveraged). "
   "This is correct: to replicate the certificate's payoff profile near B, you need more than one "
   "share of the stock funded by short bond.")

qa("Could you have used a different pricing model for Part 2?",
   "Yes — alternatives include: (1) Closed-form formula for continuous barrier options "
   "(Rubinstein & Reiner 1991) — faster but assumes constant vol and ignores discrete monitoring; "
   "(2) Monte Carlo simulation — flexible but slower and noisier for Greeks; "
   "(3) Heston stochastic vol model — more realistic vol dynamics but requires additional calibration. "
   "The binomial tree was chosen because it handles path dependence naturally, produces Greeks via "
   "central differences at the same cost, and is transparent enough to inspect node-by-node.")

qa("What is Brent's method and why use it for Task XI?",
   "Brent's method (scipy.optimize.brentq) is a bracketed root-finding algorithm. Given f(a) < 0 "
   "and f(b) > 0, it guarantees convergence to a root in [a,b]. It combines bisection (always "
   "stays in bracket), secant method (fast near the root), and inverse quadratic interpolation. "
   "We use it because VaR(α) is monotone in α (confirmed by preliminary scan), so there is exactly "
   "one α* where VaR(α*) = −15%. No analytical formula exists, so numerical root-finding is required.")

qa("What happens to α* if drift is lower?",
   "If μ = 8% instead of 54%, the 5th percentile of terminal stock prices falls below S₀. "
   "Puts activate at the VaR quantile, reducing VaR. VaR(α) becomes decreasing in α, the "
   "VaR-vs-α curve slopes downward, and α* becomes the minimum allocation that provides "
   "VaR ≤ 15% (i.e., a meaningful protection target). The regulatory constraint becomes "
   "a floor, not a ceiling. This is the economically normal regime.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# DATA SOURCES
# ══════════════════════════════════════════════════════════════════════════════

heading("Data Sources Reference", level=1, color=(0, 79, 159))

rows_data = [
    ("RHM.DE daily prices (2023–2026)", "Yahoo Finance via yfinance, auto_adjust=True"),
    ("DU2076 market closing prices", "onvista.de — DU2076 historical prices (Schluss column)"),
    ("DU2076 product terms (K, B, maturity)", "DZ BANK — dzbank-wertpapiere.de/DU2076 + KID PDF"),
    ("ECB Svensson yield curve params", "ECB Data Portal — AAA yield curve parameters (daily)"),
    ("Rheinmetall FY2025 dividend", "ir.rheinmetall.com/investor-relations/share/dividend"),
    ("DDV market statistics", "derivateverband.de/DEU/Statistiken/Marktentwicklung"),
    ("Stress scenario magnitudes (FRTB)", "BCBS Minimum capital requirements for market risk, Jan 2019, §21.83"),
    ("Stress scenario magnitudes (PRIIPs)", "EU Regulation 1286/2014; ESMA/2017/1253"),
]

table = doc.add_table(rows=1+len(rows_data), cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Item"
hdr_cells[1].text = "Source"
for h in hdr_cells:
    for run in h.paragraphs[0].runs:
        run.bold = True

for i, (item, source) in enumerate(rows_data):
    row_cells = table.rows[i+1].cells
    row_cells[0].text = item
    row_cells[1].text = source

doc.add_paragraph()

# ── Final note ────────────────────────────────────────────────────────────────
heading("Final Notes", level=2)

body(
    "All code is in B401_Assignment_2026.ipynb. The notebook is fully self-contained — "
    "data is either downloaded at runtime (Yahoo Finance) or loaded from local CSV files "
    "(cert_prices.csv, ECB Data Portal export). All intermediate results are printed inline. "
    "The binomial tree for N=200 is also exported to binomial_tree_200.csv for inspection.",
    size=11
)
body(
    "Graphs are saved to graphs/ and embedded in the poster. All formulas in this document "
    "match the exact implementations in the notebook — no approximations or simplifications "
    "were introduced for the explanation.",
    size=11
)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = "/Users/daanishmuzaffar/Daanish/MEF/Contionuous Time derivative Pricing/Assignment/B401_Assignment_Explanation.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
