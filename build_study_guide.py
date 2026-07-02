"""
Builds B401_Study_Guide.docx — a detailed task-by-task study guide for the
B401 take-home assignment (notebook B401_Assignment_2026.ipynb).

Content is compiled from:
  - CTDP_Assignment_2026.pdf (official brief, condensed in assignment_instructions.md)
  - CTDP2026_Session_notes_combined.pdf (Sessions 2, 3 and the Q&A slides)
  - the actual notebook implementation and its executed outputs

Run:  python3 build_study_guide.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re

ACCENT = RGBColor(0x8B, 0x1A, 0x1A)   # dark red, Tübingen-ish
GREY   = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# ── base styles ───────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)

for name, size, color in [('Heading 1', 16, ACCENT), ('Heading 2', 13, ACCENT),
                          ('Heading 3', 11.5, RGBColor(0, 0, 0))]:
    st = doc.styles[name]
    st.font.name = 'Calibri'
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = color


def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def h3(text):
    doc.add_heading(text, level=3)

def p(text, italic=False, size=None, color=None):
    """Paragraph with **bold** inline markup support."""
    par = doc.add_paragraph()
    for i, chunk in enumerate(re.split(r'\*\*', text)):
        if not chunk:
            continue
        run = par.add_run(chunk)
        run.bold = (i % 2 == 1)
        run.italic = italic
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
    return par

def bullet(text, level=0):
    par = doc.add_paragraph(style='List Bullet' + ('' if level == 0 else ' 2'))
    for i, chunk in enumerate(re.split(r'\*\*', text)):
        if not chunk:
            continue
        run = par.add_run(chunk)
        run.bold = (i % 2 == 1)
    return par

def formula(text, note=None):
    """Centred serif formula line (plain-text / unicode math)."""
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run(text)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(11)
    par.paragraph_format.space_before = Pt(4)
    par.paragraph_format.space_after = Pt(4)
    if note:
        n = doc.add_paragraph()
        n.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = n.add_run(note)
        r.font.size = Pt(9)
        r.font.color.rgb = GREY
        r.italic = True

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, htxt in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(htxt)
        run.bold = True
        run.font.size = Pt(9.5)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def qa(q, a):
    par = doc.add_paragraph()
    r = par.add_run('Q: ' + q)
    r.bold = True
    par2 = p('A: ' + a)
    par2.paragraph_format.space_after = Pt(10)


# ══════════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════════
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('B401 Continuous-Time Derivatives Pricing\nTake-Home Assignment — Complete Study Guide')
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Task-by-task explanation of what was done, the theory behind it, every formula, '
                'all data sources and the full methodology\n'
                'Notebook: B401_Assignment_2026.ipynb  |  Summer Term 2026  |  '
                'Eberhard Karls Universität Tübingen (Prof. Dr. Christian Koziol / Benjamin Harsch)')
r.font.size = Pt(10.5)
r.font.color.rgb = GREY

p('**How to use this guide.** Each task section answers four questions: (1) What does the assignment '
  'ask for? (2) What theory applies (with the lecture-slide reference where the session notes cite one)? '
  '(3) What exactly does the notebook do, step by step? (4) What came out and how do we interpret it? '
  'A formula bank, a source register and a list of likely oral-defence questions close the document. '
  'All quoted numbers are the actual executed outputs of the notebook.')

# ══════════════════════════════════════════════════════════════════════════════
# 1. ASSIGNMENT AT A GLANCE
# ══════════════════════════════════════════════════════════════════════════════
h1('1. The Assignment at a Glance')

p('The take-home assignment has **three parts and eleven tasks** (Roman numerals I–XI). The tutor Q&A '
  '(Session 2 recap slide) set the complexity bar: the product must be a genuine **structured** product, '
  'at least at the level of a discount certificate — standalone calls/puts are not sufficient. Grading is '
  'driven mainly by (i) product complexity and (ii) poster quality, then technical implementation, correct '
  'valuation methodology, sensible parameters and economic reasoning.')

table(['Part', 'Theme', 'Tasks', 'What the notebook does'],
      [['Part 1', 'Designing a certificate (fictitious product, no valuation required)',
        'I–IV', 'Capped Capital Protected Participation Note (CPN) on Rheinmetall AG'],
       ['Part 2', 'Valuation of a traded certificate vs. market prices',
        'V–VII', 'DZ BANK Bonus Cap Certificate DU2076 on Rheinmetall, priced on a CRR binomial tree'],
       ['Part 3', 'Portfolio insurance via Monte Carlo simulation',
        'VIII–XI', 'EUR 10,000 in RHM over 1 year; put-option overlay; stress tests; 15% VaR constraint']],
      widths=[0.7, 2.6, 0.7, 3.2])

p('**One underlying everywhere:** Rheinmetall AG (RHM.DE, XETRA). Part 1 designs a new product on it, '
  'Part 2 values a really traded certificate on it, Part 3 simulates a portfolio invested in it. This '
  'continuity means one data pipeline and directly comparable results across all three parts.')

p('**Deliverables:** one A3 landscape poster (min. font 11, self-explanatory, name + student ID top-left, '
  'PDF), runnable commented code, all data files, and an AI-usage disclaimer. Deadline 6 July 2026, 1 pm; '
  'poster presentation 9 July 2026.')

# ══════════════════════════════════════════════════════════════════════════════
# 2. SETUP
# ══════════════════════════════════════════════════════════════════════════════
h1('2. Setup — Data, Volatility and the Risk-Free Rate')
p('Everything in the notebook rests on three pieces of shared infrastructure that are worth understanding '
  'before any individual task: the price data, the two volatility estimators, and the Svensson risk-free '
  'curve. Session 2 (slides 2–3) lists the five valuation inputs every pricing exercise needs: spot price, '
  'strike/cap/barrier, time to maturity, risk-free rate, and volatility. This section is where the notebook '
  'sources each of them.')

h2('2.1 Price data (spot input)')
bullet('**Source:** Yahoo Finance via the yfinance package, ticker RHM.DE, dividend- and split-adjusted '
       'closes (auto_adjust=True), 02 Jan 2023 – 01 Jun 2026 (866 trading days).')
bullet('**Cache-first design:** prices are stored in data/rhm_prices.csv and read from disk on every run '
       '(REFRESH = False). No live API call is made by default, so the submitted notebook runs fully '
       'offline and always reproduces the same numbers — a direct response to the requirement that the '
       'grader can re-run the code and get identical results.')
bullet('**Why start in Jan 2023:** the DU2076 valuation window opens 04 Sep 2025; ~2.7 years of history '
       'before that provides warm-up data for the rolling volatility and spans a full volatility regime '
       '(calm uptrend and the Nov 2025 crash).')
bullet('**Why end = 2026-06-02:** yfinance treats the end date as exclusive, so this setting makes '
       '01 Jun 2026 — the design date for Part 1 and the last valuation day for Part 2 — the final row.')

h2('2.2 Log-returns and the two volatility estimators')
p('All volatility estimation works on daily continuously compounded (log) returns:')
formula('rₜ = ln(Sₜ / Sₜ₋₁)')
p('The notebook deliberately maintains **two different volatility estimators** and never merges them:')
table(['Estimator', 'Formula', 'Value', 'Used for'],
      [['Full-sample historical σ (SIGMA_HIST)', 'σ̂ = √252 · std(r₁ … r_T),  T ≈ 866 days',
        '39.88% p.a.', 'Part 1 product design (SIGMA_CPN); reference through-cycle vol'],
       ['30-day rolling historical σ (rolling_vol_30)', 'σ̂ₜ = √252 · std(r₍ₜ₋₂₉₎ … rₜ), clipped to [10%, 60%]',
        'range 19.1%–77.6%, mean 37.8%', 'Part 2 daily pricing (Tasks V–VII), looked up per valuation day'],
       ['252-day trailing σ (sigma_mc)', 'σ̂ = √252 · std(last 252 returns to 01 Jun 2026)',
        '45.57% p.a.', 'Part 3 Monte Carlo diffusion vol and put-pricing vol']],
      widths=[1.9, 2.6, 1.2, 1.5])
p('**Why the split?** A stable long-run estimate is the right input when designing a product or running a '
  'one-year simulation. But for daily mark-to-model pricing, a contemporaneous risk measure matters more: '
  'the 30-day window captures the volatility regime prevailing on each valuation day instead of averaging '
  'over three years. This time-varying choice is a deliberate quality feature of Task V (the tree re-reads '
  'the current vol every day), not an inconsistency. The [0.10, 0.60] clip guards against degenerate '
  'estimates in extremely quiet or extremely turbulent 30-day windows.')

h2('2.3 Risk-free rate — Svensson (1994) yield curve')
p('Session 2 (slide 4) endorses the **Svensson approach** for the risk-free input (a constant rate is only '
  'the pass-level fallback). The Svensson model writes the continuously compounded spot rate for maturity '
  't as a six-parameter function that can reproduce all typical yield-curve shapes:')
formula('₀yₜ = β₀ + β₁·(1−e^(−t/τ₁))/(t/τ₁) + β₂·[(1−e^(−t/τ₁))/(t/τ₁) − e^(−t/τ₁)] '
        '+ β₃·[(1−e^(−t/τ₂))/(t/τ₂) − e^(−t/τ₂)]')
p('The six parameters (β₀, β₁, β₂, β₃, τ₁, τ₂) are published **daily** — the notebook uses the ECB Data '
  'Portal estimates for the AAA-rated euro-area sovereign curve, downloaded once and stored locally in '
  'data/ECB Data Portal_20260615151519.csv (199 dates, Sep 2025 – Jun 2026). Interpretation of the terms: '
  'β₀ is the long-run level, β₁ controls the short-end slope, β₂ and β₃ add up to two humps at maturities '
  'governed by τ₁ and τ₂.')
p('The notebook uses the curve at **three different points**, each deliberate:')
bullet('**Part 1 design:** R_CPN = svensson(3.0) = **2.62%** — the 3-year spot rate matching the CPN maturity, '
       'from the static 01 Jun 2026 parameter snapshot SV_P1.')
bullet('**Part 2 daily loop:** get_risk_free_rate(date, T_rem) looks up **that day\'s** ECB parameters '
       '(forward-filling weekends/holidays) and evaluates the curve at the certificate\'s **remaining '
       'maturity** on that day — a daily, maturity-matched discount rate (mean 2.14% over the window, '
       'range 1.87%–2.59%).')
bullet('**Parts 1 & 3 benchmark:** RISK_FREE = svensson(10.0) = **3.07%**, the 10-year benchmark rate used '
       'as the Sharpe-ratio reference and the Part 3 option-pricing rate. Using the long benchmark here is '
       'an intentional design choice.')

h2('2.4 Dividend yield')
p('Rheinmetall\'s FY 2025 dividend is **EUR 11.50 per share** (source: Rheinmetall IR, ex-div 13 May 2026). '
  'Divided by the design-date spot of EUR 1,207 this gives a continuous dividend-yield input q = **0.95%**. '
  'Session 3 confirms dividends should be included where relevant: in Black-Scholes the spot is replaced by '
  'S₀·e^(−qT); in the tree the risk-neutral probability uses (r − q); in the MC drift q lowers the '
  'risk-neutral drift to r − q (for the physical paths the adjusted price series already reinvests dividends).')

# ══════════════════════════════════════════════════════════════════════════════
# PART 1
# ══════════════════════════════════════════════════════════════════════════════
h1('3. PART 1 — Designing the Certificate (Tasks I–IV)')
p('Part 1 designs a **fictitious** product. The Q&A "Do\'s/Don\'ts" slide is explicit: define an investor '
  'profile, design an appropriate certificate, describe the components, illustrate the payoff, assess the '
  'market volume — and **no valuation is required, no unrealistic components**. Any pricing that appears in '
  'Part 1 is only an internal sanity check on the participation rate.')

h2('3.1 Task I — Investor Profile')
h3('What was asked')
p('Define a specific, realistic investor (age, horizon, risk tolerance, income vs. growth, market '
  'expectation) and argue why the designed product **beats a direct investment in the underlying for this '
  'investor**.')
h3('What we did')
p('The profile is a German retail investor, **age 40–60, pre-retirement**, with a 3–5 year horizon and a '
  'strong thematic conviction on the European defence re-rating (Germany\'s €100bn Sondervermögen, NATO '
  'spending, Rheinmetall\'s ~€40bn order backlog) — but **low-to-moderate risk tolerance**: at ~40% '
  'annualised volatility, Rheinmetall can plausibly lose 30–40% in a year, which is incompatible with a '
  'capital-preservation mindset. Expectations: structurally bullish, central scenario +15–30% over 3 years; '
  'nominal capital protection at maturity is non-negotiable.')
h3('Why the product beats the alternatives (the argument that earns the marks)')
table(['Alternative', 'Why it fails this investor'],
      [['Direct RHM stock', 'Right theme, but ~40% vol and full single-stock drawdown risk — too much downside'],
       ['Defence ETF', 'Diversifies away the Rheinmetall-specific alpha; still no capital floor'],
       ['Tagesgeld / fixed income', 'No access to the defence theme; declining real yields'],
       ['Reverse convertible on RHM', 'Embedded short put hands the investor exactly the downside they want to avoid'],
       ['The designed CPN', 'ZCB floor removes nominal loss risk; ~85% participation captures the 15–30% central case; cap at 130% is a cheap concession for this profile']],
      widths=[1.9, 5.2])

h2('3.2 Task II — Product Design')
h3('What was asked')
p('A marketable certificate with **at least two components**; explain the intuition and describe each '
  'component. The theory source is the financial-engineering toolbox of Session 1 (slides 13–21): retail '
  'structured products are assembled from the underlying, vanilla calls/puts, binary options and barrier '
  'options (e.g. discount certificate = underlying − short call; bonus certificate = underlying + '
  'down-and-out put; capital protection = zero bond + call).')
h3('The product: Capped Capital Protected Participation Note (CPN) on RHM.DE')
p('**Intuition.** Put most of the capital into a zero-coupon bond that guarantees the nominal at maturity. '
  'The residual "option budget" buys an at-the-money call for upside participation. Selling an '
  'out-of-the-money call at 130% of S₀ brings in extra premium that lifts the participation rate α above '
  'what the budget alone would buy. The investor gives up gains beyond +30% — a mild concession when the '
  'central scenario is +15–30%. A useful engineering property at Rheinmetall\'s high σ: both calls in the '
  'spread gain from volatility, so the **spread price — and hence α — is fairly stable in σ**.')
table(['Component', 'Position', 'Role'],
      [['Zero-coupon bond, face = S₀ = €1,207', 'Long', 'Guarantees 100% of nominal at maturity'],
       ['ATM European call, K = S₀ = €1,207', 'Long α units', 'Participation in upside above S₀'],
       ['OTM European call, K = 130%·S₀ = €1,569', 'Short α units', 'Caps the gain at ~+25.6%; its premium funds a higher α']],
      widths=[2.6, 1.1, 3.4])
h3('The at-par condition and the participation rate')
p('The product is issued at par: its components must cost exactly the S₀ = €1,207 the investor pays. '
  'That single equation pins down α:')
formula('CPN₀ = ZCB₀ + α · [Call(K = S₀) − Call(K = Cap)] = S₀')
formula('α = (S₀ − ZCB₀) / [Call(S₀) − Call(Cap)]')
p('The calls are valued with the dividend-adjusted **Black-Scholes** formula (Session 2, slide 58):')
formula('C₀ = S₀·e^(−qT)·N(d₁) − K·e^(−rT)·N(d₂)')
formula('d₁,₂ = [ln(S₀/K) + (r − q ± ½σ²)·T] / (σ·√T)')
p('Inputs: r = 2.62% (3y Svensson), σ = 39.88% (full-sample historical), q = 0.95%, T = 3y. Executed result:')
table(['Quantity', 'Value (EUR)'],
      [['ZCB present value  S₀·e^(−rT)', '1,115.79'],
       ['ATM call (K = 1,207)', '338.64'],
       ['Cap call (K = 1,569), short', '231.59'],
       ['Spread (long − short)', '107.04'],
       ['Participation rate α = (1,207 − 1,115.79) / 107.04', '85.2%'],
       ['Maximum gain α·(Cap − K) → max payoff', '+25.6% (payoff €1,515.53)']],
      widths=[4.5, 2.0])
p('**Note this is not a "valuation"** in the Part 2 sense: Part 1 requires none. The Black-Scholes call '
  'here only verifies that the advertised participation rate is economically consistent with market inputs.')

h2('3.3 Task III — Payoff Profile')
h3('What was asked')
p('One clear graphic of the payoff at maturity versus the underlying, explicitly marking where the product '
  'does and does not outperform a direct investment.')
h3('The payoff function and its three regimes')
formula('CPN_T = S₀ + α · [ max(S_T − S₀, 0) − max(S_T − Cap, 0) ]')
table(['Terminal region', 'Payoff', 'vs. direct stock'],
      [['S_T ≤ 1,207  (bearish)', 'Flat €1,207 — capital protected', 'CPN outperforms (stock loses, CPN does not)'],
       ['1,207 < S_T ≤ 1,569  (moderately bullish)', '€1,207 + 0.852·(S_T − 1,207)', 'Stock slightly ahead (α < 1), CPN close behind'],
       ['S_T > 1,569  (strongly bullish)', 'Capped at €1,515.53 (+25.6%)', 'Stock outperforms — upside beyond +30% is foregone']],
      widths=[2.3, 2.6, 2.6])
p('The notebook plot (graphs/payoff_profile.png) overlays the CPN, the direct stock and a risk-free bond '
  '(which grows to ≈ €1,306), draws vertical markers at the protection level, the cap and the **break-even '
  'point vs. the bond** at S_T = K + (bond value − S₀)/α ≈ €1,323, and shades the outperformance and '
  'underperformance regions — exactly the "mark where it does and does not outperform" requirement. A '
  'table beneath quantifies CPN, stock and bond values at representative terminal prices.')

h2('3.4 Task IV — Market Size Estimation')
h3('What was asked')
p('Estimate the total German market for the product **class** using comparable products and investor demand. '
  'The anchor is the lecture\'s market table (BSW/DDV 2023): total German structured-product market '
  '≈ **EUR 112 bn** outstanding.')
h3('The funnel calculation')
table(['Step', 'Value', 'Source / logic'],
      [['Total German structured-product market', 'EUR 112 bn', 'BSW 2023 table (lecture, slide 13)'],
       ['Kapitalschutz (capital-protection) segment, 3.8%', 'EUR 4.26 bn', 'DDV product-category share'],
       ['of which participation notes (~55%)', 'EUR 2.34 bn', 'DDV sub-segment split'],
       ['Active issues ≈ 130', '≈ 6 issuers × ~20 products', 'DZ BANK, Commerzbank, LBBW, Goldman, HVB, SocGen'],
       ['Average issue size', '≈ EUR 33 mn', '4.26 bn ÷ 130 — plausible size for a RHM-linked CPN on EUWAX']],
      widths=[3.0, 1.7, 2.8])
p('**Demand reasoning** (the qualitative half of the answer): Germany\'s large cohort of '
  'security-oriented savers with sectoral convictions; declining deposit rates making capital-protected '
  'equity alternatives attractive vs. Tagesgeld; rising equity participation among 40–65-year-olds mainly '
  'through loss-protected instruments; and MiFID II suitability — capital-protected notes qualify for '
  'conservative profiles. Complexity-wise the 3-component CPN sits between discount certificates '
  '(2 components) and bonus certificates.')

# ══════════════════════════════════════════════════════════════════════════════
# PART 2
# ══════════════════════════════════════════════════════════════════════════════
h1('4. PART 2 — Valuation of a Traded Certificate (Tasks V–VII)')

h2('4.1 The product: DU2076 Bonus Cap Certificate (DZ BANK)')
p('Part 2 requires a **really traded** product on the Part 1 underlying with at least one year (≥100 '
  'trading days) of daily prices — no open-end products. The notebook uses DZ BANK\'s Bonus Cap '
  'Certificate **DU2076 / DE000DU20767** on Rheinmetall:')
table(['Field', 'Value'],
      [['Issuer / type', 'DZ BANK AG — Bonus Cap Certificate (Bonus-Zertifikat mit Cap)'],
       ['Underlying', 'Rheinmetall AG (DE0007030009, XETRA)'],
       ['Knock-out barrier B (continuous)', 'EUR 1,050 — observed daily 04 Sep 2025 → 18 Jun 2027'],
       ['Bonus / cap level K', 'EUR 2,000, ratio 1:1'],
       ['Issue / last valuation / payment', '04 Sep 2025 / 18 Jun 2027 / 25 Jun 2027'],
       ['Data', 'Product page dzbank-wertpapiere.de/DU2076; KID PDF (bib-service.dzbank.de); daily closes exported from onvista.de into data/cert_prices.csv']],
      widths=[2.6, 4.5])
p('**Payoff at maturity:**')
formula('V_T = 2,000                    if S_t > 1,050 for ALL t in [04 Sep 2025, 18 Jun 2027]')
formula('V_T = min(S_T, 2,000)      if the barrier was touched at any time')
p('**Intuition:** as long as Rheinmetall never trades at or below €1,050 during the product\'s life, the '
  'holder receives the full €2,000 bonus no matter where the stock ends. One touch of the barrier and the '
  'bonus is permanently lost — the certificate degrades into a capped long stock position. The holder '
  'gives up dividends in exchange for conditional downside protection plus a bonus above the current spot.')

h2('4.2 The pricing engine — decomposition and the binomial tree')
h3('Decomposition')
p('The certificate is split into two building blocks that are priced simultaneously on one tree:')
formula('V_cert(S₀) = V_DO-put(S₀; K = 2,000, B = 1,050) + V_capped-fwd(S₀; K = 2,000)')
bullet('**Capped forward:** pays min(S_T, K) at maturity — economically "own the stock, sell a call at K".')
bullet('**Down-and-out put:** pays max(K − S_T, 0) **only if** the stock never touched B. Together, in the '
       'no-touch state they sum to min(S_T,K) + (K − S_T)⁺ = K = the €2,000 bonus; after a knock-out only '
       'min(S_T, K) survives. The decomposition reproduces the term sheet exactly.')
h3('Why Black-Scholes cannot price this (the "path dependency" argument)')
p('Black-Scholes prices payoffs that depend **only on the terminal price S_T** — it models the lognormal '
  'distribution of the endpoint and knows nothing about the path in between. The DU2076 barrier condition '
  'is a statement about the **whole path**: touching €1,050 in March 2026 kills the bonus even if the '
  'stock recovers to €1,900 by maturity. A model with no intermediate time steps cannot answer "did the '
  'stock ever cross the barrier?". Any such product (barriers, Asians, lookbacks) needs a lattice, Monte '
  'Carlo, or a specifically derived closed form. Hence the **CRR binomial tree**.')
h3('CRR tree mechanics (Session 2, slides 57–59)')
p('Time to maturity is discretised into N steps of length Δt = T/N. Each step the stock moves up by u or '
  'down by d, calibrated to the volatility:')
formula('u = e^(σ√Δt),   d = 1/u = e^(−σ√Δt)')
p('Because u·d = 1 the tree **recombines**: at step n there are only n+1 nodes, with node (n, j) at '
  'S₀·uʲ·d^(n−j). Under a continuous dividend yield θ, imposing no-arbitrage on one step '
  '(S₀ = e^(−rΔt)·E[S_next·e^(θΔt)]) gives the **risk-neutral up-probability** (slide 59):')
formula('q = (e^((r−θ)Δt) − d) / (u − d)')
p('Note q is *not* the real-world up-probability — the arbitrage-free price needs only the risk '
  'characteristics (u, d, i.e. σ), never the true drift. **Backward induction** then proceeds:')
bullet('1. Terminal payoffs at step N: V_DO = max(K − S,0) zeroed where S ≤ B; V_cf = min(S, K).')
bullet('2. Roll back one step: V(n,j) = e^(−rΔt)·[q·V(n+1, j+1) + (1−q)·V(n+1, j)].')
bullet('3. **Barrier absorption at every step:** any node with S ≤ B = 1,050 has its down-and-out-put value '
       'forced to zero before rolling back further. Knocked-out value propagates correctly to the root.')
bullet('4. The root value V(0,0) is the model price.')
p('**Implementation details worth citing in a defence:** the pricer averages an N-step and an (N+1)-step '
  'tree — barrier-option tree prices oscillate between even and odd N because the barrier falls between '
  'node layers; averaging cancels this odd-even oscillation and smooths the Greeks. The daily valuation '
  'loop runs N = 200 steps, which on a ~1.8-year maturity puts the step size at ~2.3 trading days.')

h2('4.3 Task V — Daily Valuation over 100+ Days')
h3('What was asked')
p('Value the product with the preferred method over a window of at least 100 trading days, compare daily '
  'model prices with observed market prices, report error metrics **in a table** and comment on why model '
  'and market differ.')
h3('Methodology — the daily loop')
p('Market prices: 183 daily DU2076 closes (04 Sep 2025 – 01 Jun 2026, the full trading history of the '
  'certificate) exported from onvista.de (semicolon-separated, German decimal format) into '
  'data/cert_prices.csv and aligned with the stock\'s trading dates. Then, **for each trading day**:')
bullet('1. Remaining maturity T_rem = (18 Jun 2027 − date)/365 — recomputed daily, as the brief requires.')
bullet('2. Risk-free rate r(T_rem) from **that day\'s** ECB Svensson parameters (maturity-matched, '
       'forward-filled over non-business days).')
bullet('3. Volatility σₜ = 30-day rolling historical vol as of that day (clipped to [10%, 60%]); over the '
       'window it ranged 26.3%–58.6%, mean 44.4%.')
bullet('4. Model price = bonus_cap_cert_price(Sₜ, 2000, 1050, rₜ, q = 0.95%, σₜ, T_rem) on the N/N+1 tree.')
h3('Error metrics — definitions and results')
formula('ME = (1/n)·Σ(P̂ᵢ − Pᵢ)        MAE = (1/n)·Σ|P̂ᵢ − Pᵢ|')
formula('RMSE = √[(1/n)·Σ(P̂ᵢ − Pᵢ)²]      MAPE = (1/n)·Σ|P̂ᵢ − Pᵢ|/Pᵢ · 100')
p('ME reveals the **direction** of the bias (over- vs. under-pricing); MAE and RMSE the magnitude, with '
  'RMSE penalising large misses; MAPE scales by the price level. Executed results over 183 days:')
table(['Metric', 'Value', 'Reading'],
      [['Mean Error (ME)', '−94.25 EUR', 'Systematic UNDER-pricing by the model'],
       ['MAE', '110.89 EUR', '≈ 7% of a ~EUR 1,500 certificate'],
       ['RMSE', '123.31 EUR', 'Errors fairly uniform; no huge outliers dominating'],
       ['MAPE', '7.11%', 'Relative scale of the misfit'],
       ['Error quantiles (25/50/75/95%)', '−153.6 / −95.0 / −56.4 / +75.6 EUR', 'Bulk of days below zero — bias, not noise']],
      widths=[2.2, 2.0, 3.3])
h3('Why is there a pricing error? (the marks-earning discussion)')
bullet('**Issuer margin:** every observed onvista close embeds DZ BANK\'s manufacturing and distribution '
       'spread over fair value; a risk-neutral tree cannot reproduce that, so the model sits below the '
       'quote roughly by the spread.')
bullet('**Implied vs. realized volatility (Session 2, slides 5–7):** the market prices on *implied* vol '
       '(expectations, risk premium, supply/demand); our tree uses *realized* (historical) vol. On average '
       'implied > realized — options are on average "too expensive", so option-short structures like '
       'discount certificates look cheap. For DU2076 the twist is that the certificate has **negative '
       'vega** (higher vol → higher barrier-breach probability → smaller bonus expectation). The 30-day '
       'realized vol (mean 44.4%) exceeded what the market was implicitly pricing, so the model assumed '
       '*too much* knock-out risk and systematically underpriced.')
bullet('**Where the fit was best/worst:** best on 18 Sep 2025 (error ≈ −3 EUR) when rolling vol dipped to '
       '~35% and the vol gap closed; worst on 05 Jan 2026 (−245 EUR) during the high-vol turn of the year. '
       'This pattern is itself evidence for the volatility-gap explanation.')

h2('4.4 Task VI — Sensitivity Analysis (Greeks)')
h3('What was asked')
p('Plot the Greeks (at least Delta, Gamma, Vega — which must also appear on the poster; Theta/Rho may be '
  'left to the oral) and discuss. Inputs may be observed or sensible fictitious values.')
h3('Theory — analytic vs. numerical Greeks')
p('For vanilla options closed forms exist (Session 2, slide 12), with N′(x) = e^(−x²/2)/√(2π):')
formula('Δ_call = N(d₁),   Δ_put = N(d₁) − 1,   Γ = N′(d₁)/(S·σ·√T),   Vega = S·N′(d₁)·√T')
p('The bonus-cap certificate has **no closed-form Greeks** (path-dependent barrier), so the notebook uses '
  'the recommended numerical route — **central finite differences**: bump one input, re-price on the tree, '
  'difference:')
formula('Δ ≈ [V(S+h) − V(S−h)] / 2h        Γ ≈ [V(S+h) − 2V(S) + V(S−h)] / h²')
formula('Vega ≈ [V(σ+h) − V(σ−h)] / 2h')
p('**Bump-size choices (and why they matter):** h must be small relative to S or σ but large enough that '
  'the difference is not swamped by tree discretisation noise. The notebook uses a 1% spot bump for Delta, '
  'a **5% spot bump for Gamma** — the tree\'s node spacing near the reference spot is ~EUR 39, so a 1% '
  '(~EUR 14) bump is *sub-grid* and the second difference would collapse to float noise — and a 2pp vol '
  'bump for Vega. A light Gaussian smoothing (σ = 2 points) removes the residual node-crossing sawtooth, '
  'directly implementing the Q&A guidance "more data points, appropriate scaling, economic explanation" '
  'for non-smooth Greek plots.')
h3('Results at the reference point (S = 1,708, T = 1.79y, σ = 44.7% median rolling vol)')
table(['Greek', 'Value', 'Economic reading'],
      [['Price', '1,452.87 EUR', 'Model value at the reference inputs'],
       ['Delta', '+0.60', 'Certificate moves ~0.60 EUR per 1 EUR stock move — between bond-like and stock-like'],
       ['Gamma', '−0.00043', 'Negative near the cap: the embedded short call flattens Delta as S rises'],
       ['Vega', '−14.37 per 1pp vol', 'NEGATIVE — the signature Greek of this product (see below)']],
      widths=[1.3, 1.7, 4.5])
bullet('**Why negative Vega (two reinforcing channels):** (1) the capped forward is short a call at '
       'K = 2,000 — higher σ raises that call\'s value against the holder; (2) higher σ raises the '
       'probability of touching B = 1,050, destroying the bonus. Both point the same way: the holder '
       'wants LOW realised volatility.')
bullet('**Shape near the barrier:** Gamma spikes and Delta can exceed 1 as S approaches B — the knock-out '
       'probability reacts violently to spot there, making the value function highly convex. Near and '
       'above the cap all Greeks decay: the payoff is locked at €2,000 and the certificate behaves like '
       'a zero-coupon bond.')

h2('4.5 Task VII — Replicating Portfolio')
h3('What was asked')
p('For **each day** of the observation window, determine the replicating portfolio (stock units + bond); '
  'additionally plot the **equity fraction as a function of the underlying price at a fixed date**, and '
  'interpret the behaviour in falling markets.')
h3('Theory (Session 2, slides 26–29 and 57)')
p('Replication expresses the derivative as Δ units of stock plus a bond position. In the one-period model:')
formula('Δ = (C_u − C_d) / (S_u − S_d),      bond (credit) = 1/(1+r)^T · (Δ·S_d − C_d)')
p('In continuous Black-Scholes terms, a call = N(d₁) shares plus −N(d₂) zero bonds of face K. For the '
  'certificate the notebook takes Δ(t) from the tree (via the finite-difference Delta of Task VI) and backs '
  'out the bond leg from the identity:')
formula('V(t) = Δ(t)·S(t) + B(t)      ⇒      B(t) = V(t) − Δ(t)·S(t)')
formula('Equity fraction(t) = Δ(t)·S(t) / V(t)')
p('The equity fraction says what share of the certificate\'s value is effectively "stock". It can exceed '
  '100%: near the barrier Δ → 1 while V falls, so the replicating book is long more than one share and '
  '**short the bond** — mild leverage, exactly the hump shape shown in the lecture\'s S&P 500 example.')
h3('What the two plots show (executed results)')
bullet('**Time series (replicating_portfolio_time.png):** over Sep 2025 – Jun 2026 Rheinmetall fell from '
       '~€1,708 toward the barrier zone. Delta started ~0.25 (deep in the bonus zone: mildly-sensitive, '
       'mostly bond-financed) and climbed to ~0.9–1.0 by May 2026 — full range [0.14, 1.07]. The equity '
       'fraction rose from ~20–40% to a peak of **112%** (range 16.4%–112%; above 50% on 143 of 183 days).')
bullet('**Static snapshot (equity_fraction_vs_S.png):** sweeping S at fixed T_rem ≈ 1.79y and median vol, '
       'the fraction peaks sharply just above B = 1,050, declines through the bonus zone, and → 0 for very '
       'high S (certificate ≈ bond paying €2,000). Below B the certificate has knocked out and is modelled '
       'as a capped forward, whose equity fraction rises steeply as S falls — full downside, no floor.')
h3('The falling-market interpretation (the graded punchline)')
p('A **capital-protected** note de-risks itself in a crash (Δ → 0). A **bonus certificate does the '
  'opposite**: as S falls toward the barrier, Δ rises toward 1 and the holder ends up with near-total '
  'stock exposure at exactly the worst moment — and one barrier touch converts the product permanently '
  'into a capped forward with full downside. The replication analysis makes this hidden exposure '
  'quantitative and is the reason a holder must act well before the barrier is reached.')

# ══════════════════════════════════════════════════════════════════════════════
# PART 3
# ══════════════════════════════════════════════════════════════════════════════
h1('5. PART 3 — Portfolio Insurance Strategy (Tasks VIII–XI)')

h2('5.1 Common machinery: Monte-Carlo GBM simulation')
p('Setup fixed by the brief: initial capital **EUR 10,000**, horizon **T* = 1 year**, insurance instrument '
  '= **put options**. The stock follows Geometric Brownian Motion, dS = μS dt + σS dW, simulated at daily '
  'frequency by exact discretisation of the log-price (Session 3, slide 145):')
formula('ln S₍ᵢ₊₁₎Δt = ln SᵢΔt + (μ − ½σ²)·Δt + zᵢ₊₁·σ·√Δt,   z ~ N(0,1)')
table(['Parameter', 'Value', 'Note'],
      [['S₀', 'EUR 1,207.00', 'Last cached close, 01 Jun 2026'],
       ['μ (physical drift)', '54.05% p.a.', 'Annualised mean of daily log-returns 2023–2026 — see caveat below'],
       ['σ', '45.57% p.a.', '252-day trailing historical vol; also the put-pricing vol per the brief'],
       ['r', '3.07%', '10y Svensson benchmark (RISK_FREE)'],
       ['Paths × steps', '50,000 × 252', 'Seed fixed at 2026 → fully reproducible'],
       ['q', '0.95%', 'Enters the put pricing (risk-neutral drift r − q)']],
      widths=[1.7, 1.7, 4.1])
bullet('**Physical vs. risk-neutral drift — the conceptual heart of Part 3:** real-world return '
       'distributions (Tasks VIII–XI) are simulated under the **physical** drift μ; option **prices** '
       'always use the **risk-neutral** drift r − q inside Black-Scholes. Mixing them up is the classic '
       'error the session notes warn about ("Drift parameter → risk-neutral measure!" for pricing).')
bullet('**Drift caveat stated in the notebook:** μ = 54% reflects Rheinmetall\'s exceptional 2022–26 '
       'defence re-rating and is not a long-run forecast (a conventional ERP would be ~6–9% over r). All '
       'Part 3 conclusions are conditional on that regime; several signs (e.g. how VaR responds to more '
       'insurance) would flip under a normal drift. Stating this explicitly is part of the "reasonable '
       'parameter choice" grading criterion.')
bullet('**One shared path set** for all of Tasks VIII–XI: differences between strategies are then purely '
       'strategy effects, never Monte-Carlo noise.')
bullet('**Why no −½σ² correction in the notebook\'s step:** the drift is estimated as the mean of *log* '
       'returns, so it already is the drift of ln S; adding the usual −½σ² would double-count. (The −½σ² '
       'form applies when μ is an arithmetic/expected-price drift, as in risk-neutral pricing where '
       'E[S_T] = S₀·e^(rT) must hold.)')

h2('5.2 Task VIII — Performance Without Risk Management')
h3('What was asked')
p('Simulate the 1-year payoff distribution of EUR 10,000 fully invested in the stock; report meaningful '
  'performance **and** risk measures (the session notes explicitly name performance, risk, and '
  '"descriptive measures" like skewness and kurtosis) and explain them.')
h3('The measures (definitions used by performance_stats)')
bullet('**Mean / median / std of log-returns** ln(S_T/S₀) — location and dispersion.')
bullet('**Skewness** — asymmetry of the distribution (0 for a normal); **excess kurtosis** — tail '
       'heaviness beyond the normal (0 for a normal).')
bullet('**Sharpe ratio** = (mean − r)/std — return per unit of total risk.')
bullet('**95% VaR** = −(5th percentile of simulated returns): the loss that is not exceeded with 95% '
       'confidence over the year. **95% CVaR** = −mean of returns at or below that quantile: the expected '
       'loss *given* the tail is hit — always ≥ VaR and more informative about tail shape.')
bullet('**P(loss)** = share of paths ending below the initial wealth.')
h3('Executed results — the unhedged baseline')
table(['Measure', 'Value'],
      [['Mean / median return', '54.03% / 54.18%'],
       ['Std deviation', '45.61%'],
       ['Skewness / excess kurtosis', '0.007 / 0.048  (log-returns ≈ normal, by construction of GBM)'],
       ['Sharpe ratio', '1.117'],
       ['95% VaR / 95% CVaR', '20.64% / 40.14%'],
       ['P(loss)', '11.7%']],
      widths=[3.0, 4.0])
p('Reading: the enormous historical drift makes the average outcome spectacular, but the tail is fat in '
  'euro terms — one year in twenty loses more than a fifth of the capital, and conditional on landing '
  'there the average loss is 40%. This baseline VaR of **20.6% > 15%** is what forces genuine hedging '
  'decisions in Tasks IX–XI. Log-returns under GBM are normal, so skew/kurtosis ≈ 0 here — they become '
  'interesting once the put overlay deforms the distribution.')

h2('5.3 Task IX — Portfolio Insurance with Put Options')
h3('What was asked')
p('Add put options to hedge the downside; analyse how (i) the **fraction α of initial wealth** spent on '
  'puts and (ii) the **put strike K** affect performance and risk. The puts are **fictitious** and must be '
  'priced at the **historical volatility** of the underlying (Session 3, Task 9/10 slide).')
h3('Mechanics of insured_return')
bullet('Put premium P₀ = Black-Scholes put(S₀, K, r, q, σ_hist = 45.57%, T = 1).')
bullet('n_puts = α·W₀ / P₀ bought; n_shares = (1−α)·W₀ / S₀ into the stock.')
bullet('Terminal wealth W_T = n_puts·max(K − S_T, 0) + n_shares·S_T evaluated on all 50,000 paths; '
       'returns are ln(W_T/W₀).')
bullet('Grid: α ∈ {5, 10, 15, 20%} × K ∈ {90, 95, 100, 105% of S₀} — 16 combinations spanning the '
       'retail-realistic range (grid choice anchored to Rubinstein & Leland 1981 / Leland 1985).')
h3('Executed results (key rows of the 4×4 grid)')
table(['α', 'K', 'Mean', 'Std', 'Sharpe', '95% VaR', '95% CVaR', 'P(loss)'],
      [['0% (unhedged)', '—', '54.0%', '45.6%', '1.117', '20.6%', '40.1%', '11.7%'],
       ['5%', '105%', '49.9%', '43.7%', '1.073', '18.2%', '30.7%', '14.1%'],
       ['10%', '95%', '45.2%', '42.4%', '0.994', '18.7%', '21.7%', '17.0%'],
       ['15%', '105%', '40.8%', '40.6%', '0.929', '13.6%', '14.8%', '20.4%'],
       ['20%', '105%', '35.7%', '39.3%', '0.829', '15.1%', '16.3%', '24.4%']],
      widths=[1.1, 0.8, 0.9, 0.9, 0.9, 0.95, 1.0, 0.9])
h3('The three patterns to understand (and defend)')
bullet('**VaR is U-shaped in α:** a moderate put sleeve cuts the 5th-percentile loss (puts pay exactly in '
       'the states that define VaR), minimum around α ≈ 10–15%; beyond that the certain premium outweighs '
       'the uncertain payoff and VaR *rises again* (see α = 20% row). The optimum is interior.'),
bullet('**Higher strikes cut VaR more per euro:** a 105% (slightly ITM) put is the only one paying '
       'meaningfully at the 5th-percentile outcome of this high-drift distribution; 90% OTM puts barely '
       'move the tail. Lowest cell: VaR 13.6% at (α = 15%, K = 105%).'),
bullet('**Mean return falls monotonically in α** — premium drag — while CVaR and skewness improve: the '
       'overlay buys a *shape* change (truncated left tail, positive skew), paying with expected return. '
       'This is the classic lecture result: put fraction up ⇒ mean and vol down, skew up, VaR/CVaR '
       'sharply down.')

h2('5.4 Task X — Stress Scenario Analysis')
h3('What was asked')
p('For one specific Task IX allocation, re-price the puts with pricing volatility = historical σ ± 5pp, '
  'and/or impose a sudden −20% drop in the underlying after half a year; recompute all measures and '
  'discuss. The session notes stress that the ±5pp applies to the **pricing vol of the puts** (put prices '
  'are based on the historical vol of the underlying).')
h3('Design')
bullet('Chosen allocation: **α = 10%, K = 95%·S₀** (a middle-of-the-grid, realistic retail choice).')
bullet('Vol scenarios: σ_pricing = 45.57% ± 5pp — magnitude consistent with the FRTB/Basel vega stress '
       'convention; changes both the premium paid and (in the same run) the simulated dispersion.')
bullet('Crash scenario: all paths multiplied by 0.80 from step 126 (T/2) onward — a permanent −20% level '
       'shock in the spirit of the PRIIPs adverse scenario.')
bullet('Combined scenario: crash AND vol +5pp together — the jointly adverse "and" case of the task\'s '
       '"and/or" wording.')
h3('Executed results')
table(['Scenario', 'Mean', '95% VaR', '95% CVaR', 'P(loss)'],
      [['Baseline', '45.21%', '18.66%', '21.73%', '17.0%'],
       ['Pricing vol +5pp', '45.03%', '20.03%', '24.60%', '17.0%'],
       ['Pricing vol −5pp', '45.42%', '16.90%', '18.13%', '17.0%'],
       ['−20% shock at T/2', '25.63%', '22.35%', '24.89%', '32.2%'],
       ['Shock AND vol +5pp', '25.19%', '25.50%', '29.33%', '32.2%']],
      widths=[2.2, 1.2, 1.2, 1.2, 1.2])
p('**Interpretation:** the two stress families hit *different* channels. Vol mis-estimation works through '
  'the premium/tail-width channel (±5pp moves VaR by ±1.4–1.7pp, mean almost unchanged); the crash works '
  'through realised returns (mean −20pp, P(loss) nearly doubles to 32%), against which a 95%-strike put '
  'holding only 10% of wealth offers partial protection at best. The combined case is strictly the worst '
  '(VaR 25.5%) — the scenario a risk manager should size against. Note the baseline VaR of 18.7% already '
  'exceeds the coming 15% limit: this allocation is *not* Task XI-compliant, which motivates the search '
  'in the final task.')

h2('5.5 Task XI — Capital Requirement (VaR-Constrained Allocation)')
h3('What was asked')
p('A binding regulatory constraint: the **95% one-year VaR may not exceed 15%**. Describe the (optimal) '
  'allocation that **fully utilises** the constraint. Session notes clarifications: the requirement must '
  'bind; there are multiple valid ways to construct the portfolio; the full EUR 10,000 need not be '
  'invested; use a seed.')
h3('Methodology')
bullet('At σ = 45.57% the **unhedged VaR is 20.64% — already above the ceiling**, so unlike a low-vol '
       'world (where one would scale a compliant stock position UP to the limit) the investor here must '
       'genuinely hedge. Both Task IX levers — α and K — are searched.')
bullet('For each strike K ∈ {90, 95, 100, 105%}, VaR(α) is evaluated on a fine α-grid (0–50%, step 0.1%); '
       'the smallest α reaching the target is recorded, and among compliant (α, K) pairs the one with the '
       'highest expected return is chosen. (Monotonicity of VaR in α over the relevant range is verified '
       'first — the precondition for treating this as a clean root-finding problem, per the '
       'brentq/root-finder approach suggested in the instructions.)')
bullet('**Prudential buffer:** the VaR estimate from 50,000 paths carries ~0.1pp sampling noise; an '
       'allocation sitting exactly on 15% breaches the limit in ~45% of re-seeded runs. The notebook '
       'therefore targets 15% − 0.5pp = **14.5%**, reporting the knife-edge point too. "Fully utilise" is '
       'thus implemented as: as close to the ceiling as robust compliance allows.')
h3('Executed results')
table(['Strike K', 'α reaching 15% (edge)', 'α reaching 14.5% (safe)', 'VaR at safe α', 'E[return]'],
      [['90%', 'infeasible (VaR floors at 20.5%)', 'infeasible', '—', '—'],
       ['95%', 'infeasible (floors at 18.1%)', 'infeasible', '—', '—'],
       ['100%', 'infeasible (floors at 15.4%)', 'infeasible', '—', '—'],
       ['105%', '11.90%', '13.00%', '14.48%', '42.70%']],
      widths=[1.0, 2.1, 2.0, 1.2, 1.0])
p('**Optimal allocation: K* = 105% of S₀ (EUR 1,267), α* = 13% → EUR 1,300 in slightly-ITM 1-year puts + '
  'EUR 8,700 in the stock.** Resulting 95% VaR = 14.48% (≈0.5pp cushion under the ceiling, robust across '
  'seeds), expected return 42.70%, Sharpe 0.964. OTM and ATM puts can never reach the ceiling at any '
  'allocation — they pay too little at the 5th-percentile outcome — which is why the slightly-ITM strike '
  'is the only feasible instrument. The buffer costs ~1pp of expected return vs. the knife-edge — cheap '
  'insurance against sampling noise. More hedging than α* would be compliant but strictly return-inferior '
  '(and eventually raises VaR again via premium drag): the constraint binds, and every remaining euro of '
  'risk budget is spent on expected return.')
p('**Sensitivity worth mentioning:** on the full-sample vol (39.9%) the naked stock is already compliant '
  '(VaR ≈ 11%) and the exercise would instead scale exposure up to the limit. The notebook deliberately '
  'reports the recent-regime (45.6%) version because it forces a real hedging decision — the more '
  'conservative and more instructive reading.')

# ══════════════════════════════════════════════════════════════════════════════
# FORMULA BANK
# ══════════════════════════════════════════════════════════════════════════════
h1('6. Formula Bank (everything in one place)')

h3('Returns and volatility')
formula('rₜ = ln(Sₜ/Sₜ₋₁);    σ̂ = √252 · std(r);    rolling: std over last 30 (or 252) returns')
h3('Black-Scholes (with dividend yield q)')
formula('C = S₀e^(−qT)N(d₁) − Ke^(−rT)N(d₂);    P = Ke^(−rT)N(−d₂) − S₀e^(−qT)N(−d₁)')
formula('d₁,₂ = [ln(S₀/K) + (r − q ± ½σ²)T] / (σ√T);    put-call parity: C − P = S₀e^(−qT) − Ke^(−rT)')
h3('One-period replication (Session 2, slide 29)')
formula('Δ = (C_u − C_d)/(S_u − S_d);    C₀ = Δ·S₀ − 1/(1+r)^T · (Δ·S_d − C_d)')
h3('CRR binomial tree (with dividend yield θ)')
formula('Δt = T/N;   u = e^(σ√Δt);   d = 1/u;   q = (e^((r−θ)Δt) − d)/(u − d)')
formula('S(n,j) = S₀·uʲ·d^(n−j);    V(n,j) = e^(−rΔt)[q·V(n+1,j+1) + (1−q)·V(n+1,j)]')
formula('Barrier: V_DO(n,j) = 0 wherever S(n,j) ≤ B (checked at every step)')
h3('Greeks')
formula('Analytic: Δ_call = N(d₁);  Γ = N′(d₁)/(Sσ√T);  Vega = S·N′(d₁)·√T;  N′(x) = e^(−x²/2)/√(2π)')
formula('Numerical: Δ ≈ [V(S+h)−V(S−h)]/2h;  Γ ≈ [V(S+h)−2V(S)+V(S−h)]/h²;  Vega ≈ [V(σ+h)−V(σ−h)]/2h')
h3('Replicating portfolio / equity fraction')
formula('B(t) = V(t) − Δ(t)S(t);    Equity fraction = Δ(t)·S(t)/V(t)')
h3('GBM and Monte Carlo')
formula('dS = μS dt + σS dW;    S_T = S₀·exp[(μ − ½σ²)T + σW_T];    E[S_T] = S₀e^(μT)')
formula('Discrete: ln Sₜ₊Δt = ln Sₜ + (drift)·Δt + z·σ√Δt;   pricing drift = r − q,  real-world = μ̂')
formula('MC estimate: Ĉ₀ = e^(−rT)·(1/M)·ΣC_T⁽ᵐ⁾;   standard error = e^(−rT)·s/√M')
h3('Svensson (1994)')
formula('₀yₜ = β₀ + β₁f₁(t/τ₁) + β₂f₂(t/τ₁) + β₃f₂(t/τ₂),  fᵢ as in Section 2.3;  6 daily ECB parameters')
h3('Risk measures')
formula('95% VaR = −(5th percentile of simulated returns);    95% CVaR = −E[return | return ≤ −VaR]')
formula('Error metrics: ME = mean(P̂−P);  MAE = mean|P̂−P|;  RMSE = √mean((P̂−P)²);  MAPE in %')

# ══════════════════════════════════════════════════════════════════════════════
# SOURCES
# ══════════════════════════════════════════════════════════════════════════════
h1('7. Source Register')
table(['What', 'Source'],
      [['RHM.DE daily prices', 'Yahoo Finance via yfinance (auto-adjusted closes), cached to data/rhm_prices.csv'],
       ['Rheinmetall dividend €11.50 (FY2025)', 'Rheinmetall IR — ir.rheinmetall.com/investor-relations/share/dividend'],
       ['Svensson parameters (daily)', 'ECB Data Portal — AAA euro-area sovereign curve parameters; cached to data/ECB Data Portal_20260615151519.csv'],
       ['DU2076 terms (K = 2,000, B = 1,050, dates)', 'DZ BANK product page dzbank-wertpapiere.de/DU2076 and official KID PDF (bib-service.dzbank.de/bib/DE000DU20767-de-DE.pdf)'],
       ['DU2076 daily market closes', 'onvista.de historical-prices CSV export → data/cert_prices.csv (the data route recommended in Session 3, slide 3)'],
       ['Market-size figures', 'BSW 2023 table (lecture slide 13, total EUR 112bn) + DDV Marktvolumenstatistik (derivateverband.de) for segment shares and issuer list'],
       ['Task IX grid design', 'Rubinstein & Leland (J. Finance 1981); Leland (J. Finance 1985)'],
       ['±5pp vol stress magnitude', 'FRTB / Basel market-risk vega stress convention (BCBS Jan 2019)'],
       ['−20% crash scenario', 'PRIIPs KID adverse-scenario spirit (EU Reg. 1286/2014; ESMA 2017)'],
       ['Methodology (tree, Greeks, GBM, VaR)', 'B401 Session notes: Session 2 (slides 2–15), Session 3 (slides 3–13), Q&A slides']],
      widths=[2.8, 4.3])

# ══════════════════════════════════════════════════════════════════════════════
# LIKELY QUESTIONS
# ══════════════════════════════════════════════════════════════════════════════
h1('8. Likely Defence Questions — with Short Answers')

qa('Why can\'t Black-Scholes price DU2076?',
   'Its payoff depends on the whole price path (one touch of €1,050 permanently kills the bonus), while '
   'Black-Scholes only models the terminal distribution. Path dependency requires a tree, Monte Carlo, or '
   'a special closed form; we use a CRR tree with barrier absorption at every step.')
qa('Why do model and market prices differ in Task V?',
   'Two reasons: (1) the issuer margin embedded in every quoted price; (2) the volatility gap — the market '
   'prices on implied vol, our tree on realized (30-day rolling) vol. Because the certificate has negative '
   'vega and realized vol ran above implied, the model over-weights knock-out risk and underprices '
   '(ME = −94 EUR). The fit is best exactly when rolling vol dipped (Sep 2025) and worst in the vol spike '
   '(Jan 2026), consistent with this explanation.')
qa('Why is the certificate\'s Vega negative?',
   'Two reinforcing channels: the embedded short call at the cap loses when vol rises, and a higher vol '
   'raises the barrier-touch probability, destroying the expected bonus. The holder wants calm markets.')
qa('Why is Gamma negative near the cap and spiking near the barrier?',
   'Near the cap the payoff flattens (short-call character) so Delta falls as S rises. Near the barrier the '
   'knock-out probability reacts violently to spot, making value highly convex and pushing Delta above 1.')
qa('Why numerical Greeks, and why those bump sizes?',
   'No closed form exists for the barrier product. Central differences on the tree; the Gamma bump must '
   'exceed one tree-node spacing (~EUR 39 at N=200/500) or the second difference collapses into '
   'discretisation noise — hence 5% for Gamma vs. 1% for Delta, and 2pp for Vega.')
qa('What does an equity fraction above 100% mean?',
   'The replicating book holds more than one share financed by a short bond position — mild leverage. It '
   'happens near the barrier where Delta → 1 while the certificate value falls.')
qa('How does the bonus certificate behave in a falling market?',
   'Opposite to a capital-protected note: Delta RISES toward 1 as S approaches the barrier, so stock '
   'exposure is maximal at the worst time, and a single touch converts the product into a capped forward '
   'with full downside. That is the key risk-management lesson of Task VII.')
qa('Why do you use two different volatilities, and two different drift concepts?',
   'Vols: a stable full-sample σ for design/simulation, a 30-day rolling σ for daily mark-to-model pricing '
   '(a contemporaneous risk measure grades better than one frozen number). Drifts: real-world '
   'distributions use the physical μ̂; option prices always use the risk-neutral drift r − q — mixing them '
   'is an arbitrage error.')
qa('Why is there no −½σ² term in your MC step?',
   'Because μ̂ is estimated as the mean of log-returns, it already IS the drift of ln S. The −½σ² '
   'correction applies when converting an arithmetic price drift (like r under the pricing measure) into '
   'a log drift.')
qa('Why does more put insurance eventually RAISE VaR?',
   'The premium is a certain cost while the payoff is contingent; in a high-drift distribution the '
   '5th-percentile path often ends only mildly down, so past an interior optimum extra premium drag '
   'deepens the quantile loss. Hence the U-shape in α.')
qa('Why does Task XI need the 105% strike, and why the 0.5pp buffer?',
   'OTM/ATM puts pay too little at the 5th-percentile outcome — their VaR floors above 15% at ANY '
   'allocation. Only the slightly-ITM put moves the tail enough. The buffer: a 50,000-path VaR carries '
   'sampling noise, and a portfolio sitting exactly on the limit breaches it in ~45% of re-seeded runs; a '
   'risk manager sizes below a limit, not on it. Cost: ~1pp of expected return.')
qa('What does "fully utilising the constraint" mean if you sit at 14.48%, not 15.00%?',
   'The constraint binds economically: among robustly compliant portfolios we maximise expected return, '
   'and the chosen point is as close to the ceiling as Monte-Carlo estimation error responsibly allows. '
   'The knife-edge allocation (α = 11.9%, VaR = 15.00%) is reported alongside for transparency.')
qa('Is your Part 3 drift realistic?',
   'No, and the notebook says so: 54% p.a. is Rheinmetall\'s exceptional re-rating, not a forecast; a '
   'normal ERP would be 6–9% over the risk-free rate. Results are explicitly conditional on the regime, '
   'and the sign of some findings would change under a conventional drift.')
qa('How is reproducibility guaranteed?',
   'All external data (prices, certificate closes, ECB Svensson parameters) is cached as local files and '
   'read from disk — no live API calls at runtime; the MC seed is fixed at 2026; one shared path set '
   'serves all Part 3 tasks; and the tree averages N and N+1 steps for numerically stable prices.')

# ── footer ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
f = doc.add_paragraph()
f.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = f.add_run('Generated from B401_Assignment_2026.ipynb, the official assignment brief and the '
              'Session 2/3/Q&A notes. All figures are executed notebook outputs.')
r.font.size = Pt(8.5)
r.font.color.rgb = GREY
r.italic = True

OUT = 'B401_Study_Guide.docx'
doc.save(OUT)
print(f'Saved {OUT}')
